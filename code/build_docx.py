"""
Build the formatted Word document for Chapter 4 - Results and Discussion.

Typography requirements:
  - Times New Roman font throughout
  - Body text: 12 pt
  - Main headings (Chapter title, top-level numbered sections like 4.1, 4.2): 14 pt, bold
  - Sub-section headings (4.7.1 etc.): 12 pt, italic

Run from repository root:
    python code/build_docx.py
Outputs:
    Chapter_4_Results_and_Discussion.docx
"""

from __future__ import annotations
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BASE = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir))
FIG_DIR = os.path.join(BASE, "figures")
OUT_PATH = os.path.join(BASE, "Chapter_4_Results_and_Discussion.docx")

FONT_NAME = "Times New Roman"
BODY_PT = 12
MAIN_PT = 14
SUB_PT = 12


# --------------------------------------------------------------- helpers
def set_run_font(run, *, size_pt=BODY_PT, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    # ensure East Asian fallback is also Times New Roman (Word picks this up)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph_with_runs(doc, runs_spec, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_before=0, space_after=6, line_spacing=1.5,
                            first_line_indent=None):
    """runs_spec is a list of dicts with 'text' and optional 'bold'/'italic'/'size'."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Inches(first_line_indent)
    for spec in runs_spec:
        run = p.add_run(spec["text"])
        set_run_font(run, size_pt=spec.get("size", BODY_PT),
                     bold=spec.get("bold", False),
                     italic=spec.get("italic", False))
    return p


def add_body(doc, text, *, indent_first=True):
    return add_paragraph_with_runs(doc, [{"text": text}],
                                   first_line_indent=0.4 if indent_first else None)


def add_main_heading(doc, text, *, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size_pt=MAIN_PT, bold=True)
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size_pt=SUB_PT, italic=True, bold=False)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size_pt=11, italic=True)
    return p


def add_figure(doc, fname, caption, *, width_in=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(os.path.join(FIG_DIR, fname), width=Inches(width_in))
    add_caption(doc, caption)


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    """Add single-line borders on all sides."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tc_pr.append(borders)


def add_table(doc, header, rows, *, header_shade="D9D9D9",
              col_widths=None, font_size=11, align_numeric=True):
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # header row
    hdr_cells = table.rows[0].cells
    for i, txt in enumerate(header):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(txt))
        set_run_font(run, size_pt=font_size, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, header_shade)
        set_cell_borders(cell)

    # body rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            text = str(val)
            # alignment heuristic
            if align_numeric and c_idx > 0 and any(ch.isdigit() for ch in text):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            # bold first column on body and any **wrapped** text
            if text.startswith("**") and text.endswith("**"):
                run = p.runs[0]
                run.text = text.strip("*")
                set_run_font(run, size_pt=font_size, bold=True)
            else:
                set_run_font(run, size_pt=font_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)

    if col_widths:
        for row in table.rows:
            for c, w in enumerate(col_widths):
                row.cells[c].width = Inches(w)

    # spacer paragraph after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return table


# --------------------------------------------------------------- document
def build():
    doc = Document()

    # ---- page setup ----
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # set default style to Times New Roman 12 pt
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_PT)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    if rFonts.getparent() is None:
        rpr.append(rFonts)

    # ===================================================== title block
    # CHAPTER 4
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("CHAPTER 4")
    set_run_font(run, size_pt=MAIN_PT, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("RESULTS AND DISCUSSION")
    set_run_font(run, size_pt=MAIN_PT, bold=True)

    # ===================================================== 4.1 Introduction
    add_main_heading(doc, "4.1  Introduction")
    add_body(doc,
        "This chapter presents and interprets the empirical findings of the "
        "integrated Fuzzy Analytic Hierarchy Process (Fuzzy AHP) and Fuzzy "
        "ELECTRE I framework applied to the malta (Citrus sinensis) post-harvest "
        "supply chain in four districts of Uttarakhand. Section 4.2 reports the "
        "descriptive characteristics of the 1,300-respondent dataset and the "
        "psychometric properties of the measurement scales. Section 4.3 documents "
        "the construction of the crisp and fuzzy decision matrices. Section 4.4 "
        "reports the criterion weights derived through Buckley's Fuzzy AHP, "
        "including the consistency assessment. Section 4.5 reports the application "
        "of Fuzzy ELECTRE I to obtain the outranking structure and the final "
        "ranking of supply-chain stages. Section 4.6 reports the sensitivity "
        "analysis. Section 4.7 discusses the findings in relation to the existing "
        "literature, identifies methodological observations, and develops the "
        "policy implications. Section 4.8 summarises the chapter.")

    add_body(doc,
        "Throughout this chapter, the five supply-chain stages \u2014 Farm (A1), "
        "Village Trader (A2), Mandi (A3), Retail (A4) and Transport (A5) \u2014 "
        "are treated as alternatives, and the three dimensions \u2014 Cost (C1), "
        "Time (C2) and Quality (C3) \u2014 are treated as evaluation criteria. "
        "The criteria are benefit-type: a higher mean Likert score indicates "
        "greater perceived criticality and therefore greater priority for "
        "intervention.")

    # ===================================================== 4.2 Descriptive
    add_main_heading(doc, "4.2  Descriptive analysis of the survey data")

    add_sub_heading(doc, "4.2.1  Sample composition")
    add_body(doc,
        "The dataset comprises 1,300 stakeholder responses collected across the "
        "four malta-producing districts of Uttarakhand. Table 4.1 reports the "
        "geographic distribution of the sample and Table 4.2 reports the "
        "stakeholder-group composition.")

    add_caption(doc, "Table 4.1 \u2014 Geographic distribution of respondents (N = 1,300).")
    add_table(doc,
        ["District", "Respondents", "Share (%)"],
        [["Chamoli", "340", "26.2"],
         ["Tehri Garhwal", "333", "25.6"],
         ["Rudraprayag", "316", "24.3"],
         ["Pauri Garhwal", "311", "23.9"],
         ["**Total**", "**1,300**", "**100.0**"]],
        col_widths=[2.4, 1.4, 1.4])

    add_caption(doc, "Table 4.2 \u2014 Stakeholder-group composition.")
    add_table(doc,
        ["Stakeholder group", "Respondents", "Share (%)"],
        [["Farmers", "650", "50.0"],
         ["Wholesalers", "130", "10.0"],
         ["Transporters", "130", "10.0"],
         ["Retailers", "130", "10.0"],
         ["Village Traders", "117", "9.0"],
         ["Commission Agents", "78", "6.0"],
         ["Experts / Officials", "65", "5.0"],
         ["**Total**", "**1,300**", "**100.0**"]],
        col_widths=[2.4, 1.4, 1.4])

    add_body(doc,
        "The sample is approximately balanced across districts (each contributes "
        "23.9\u201326.2 %), and farmers \u2014 the central actors in the "
        "post-harvest chain \u2014 constitute half of the sample, with the "
        "remaining 50 % distributed across the downstream actors. This composition "
        "affords sufficient statistical power for both pooled inference and "
        "stage-by-stage cross-validation.")

    add_sub_heading(doc, "4.2.2  Reliability of measurement scales")
    add_body(doc,
        "The 82 Likert items are organised into 15 latent scales, formed by the "
        "Cartesian product of the five stages and three criteria. Table 4.3 "
        "reports Cronbach's alpha and the average inter-item correlation for each "
        "scale. All alpha values exceed the conventional 0.80 threshold for "
        "excellent reliability (Nunnally and Bernstein, 1994), with a mean of "
        "0.885 and a range of [0.820, 0.919]. The instrument is therefore highly "
        "reliable.")

    add_caption(doc, "Table 4.3 \u2014 Internal-consistency reliability of the 15 scales.")
    add_table(doc,
        ["Scale", "k (items)", "Cronbach's \u03B1", "Avg. inter-item r"],
        [["Farm \u2014 Cost", "7", "0.917", "0.614"],
         ["Farm \u2014 Time", "6", "0.897", "0.591"],
         ["Farm \u2014 Quality", "6", "0.919", "0.656"],
         ["Trader \u2014 Cost", "6", "0.899", "0.597"],
         ["Trader \u2014 Time", "5", "0.847", "0.525"],
         ["Trader \u2014 Quality", "5", "0.894", "0.628"],
         ["Mandi \u2014 Cost", "6", "0.905", "0.615"],
         ["Mandi \u2014 Time", "5", "0.861", "0.553"],
         ["Mandi \u2014 Quality", "5", "0.886", "0.610"],
         ["Retail \u2014 Cost", "6", "0.885", "0.563"],
         ["Retail \u2014 Time", "4", "0.820", "0.533"],
         ["Retail \u2014 Quality", "5", "0.875", "0.583"],
         ["Transport \u2014 Cost", "6", "0.900", "0.601"],
         ["Transport \u2014 Time", "5", "0.903", "0.651"],
         ["Transport \u2014 Quality", "5", "0.870", "0.572"],
         ["**Mean (15 scales)**", "**5.4**", "**0.885**", "**0.593**"]],
        col_widths=[2.0, 1.0, 1.4, 1.6], font_size=10)

    add_sub_heading(doc, "4.2.3  Distribution of criticality scores")
    add_body(doc,
        "Figure 4.1 displays the distribution of respondent-level mean Likert "
        "scores by stage and criterion. Across all (stage, criterion) cells, "
        "medians fall in the 3.5\u20134.0 range, indicating that respondents "
        "perceive every supply-chain stage as moderately to highly critical. The "
        "substantive question, therefore, is not whether a stage is critical, but "
        "which stage is most critical \u2014 a question that the FMCDM framework "
        "is designed to answer.")

    add_figure(doc, "fig01_likert_distribution.png",
        "Figure 4.1 \u2014 Distribution of respondent-level mean criticality "
        "scores by supply-chain stage and criterion. Boxes denote the "
        "interquartile range; whiskers extend to 1.5 \u00D7 IQR; dots mark "
        "statistical outliers.")

    add_body(doc,
        "Figure 4.2 presents the Pearson-correlation heatmap among the 15 "
        "stage-by-criterion scale scores. Within-stage correlations (e.g., "
        "Farm-Cost correlated with Farm-Time) are systematically larger than "
        "across-stage correlations, supporting the latent-factor structure of the "
        "survey instrument and justifying the treatment of the three dimensions "
        "as analytically distinct criteria.")

    add_figure(doc, "fig15_scale_correlation_heatmap.png",
        "Figure 4.2 \u2014 Pearson correlation among the 15 stage-by-criterion "
        "scale scores.")

    # ===================================================== 4.3 Decision matrix
    add_main_heading(doc, "4.3  Construction of the fuzzy decision matrix")

    add_sub_heading(doc, "4.3.1  Crisp decision matrix")
    add_body(doc,
        "The 5 \u00D7 3 crisp decision matrix M is constructed by computing, for "
        "each (stage i, criterion j), the mean across all 1,300 respondents of "
        "the per-respondent mean of the items belonging to scale (i, j). Table "
        "4.4 reports the resulting matrix and Figure 4.3 visualises it as a "
        "heatmap.")

    add_caption(doc, "Table 4.4 \u2014 Crisp decision matrix M (mean Likert criticality, 1\u20135 scale).")
    add_table(doc,
        ["Stage", "Cost", "Time", "Quality"],
        [["Farm", "3.7534", "3.7755", "3.7719"],
         ["Village Trader", "3.6478", "3.6532", "3.6520"],
         ["Mandi", "3.7251", "3.7734", "3.7335"],
         ["Retail", "3.5338", "3.5438", "3.7394"],
         ["Transport", "3.6294", "3.8525", "3.6274"]],
        col_widths=[1.8, 1.3, 1.3, 1.3])

    add_figure(doc, "fig02_decision_matrix_heatmap.png",
        "Figure 4.3 \u2014 Crisp decision matrix heatmap (stage \u00D7 criterion "
        "mean criticality).")

    add_body(doc,
        "Two patterns are immediately visible in Table 4.4 and Figure 4.3. "
        "First, the Farm and Mandi stages dominate on Cost and Quality. Farm "
        "scores the highest on Cost (3.7534) and Quality (3.7719), while Mandi "
        "is a near match (3.7251 on Cost, 3.7335 on Quality). These two stages "
        "are the only ones that score above 3.70 on more than one criterion. "
        "Second, Time criticality is concentrated at Transport. Transport posts "
        "the highest Time score (3.8525), reflecting the well-documented "
        "sensitivity of transit duration in mountainous routes, but it scores "
        "comparatively lower on Cost (3.6294) and Quality (3.6274). Third, "
        "Retail is the cost-leader stage (lowest Cost, 3.5338) and the lowest on "
        "Time (3.5438), although it remains competitive on Quality (3.7394) due "
        "to consumer-facing freshness expectations.")

    add_body(doc,
        "The radar plot in Figure 4.4 makes these inter-stage profile "
        "differences explicit.")

    add_figure(doc, "fig14_radar_alternatives.png",
        "Figure 4.4 \u2014 Stage profiles in criterion space (radar plot, crisp "
        "means).")

    add_sub_heading(doc, "4.3.2  Fuzzification of the decision matrix")
    add_body(doc,
        "To preserve the empirical uncertainty in the survey responses, each "
        "crisp entry m(i, j) is fuzzified into a triangular fuzzy number "
        "M-tilde(i, j) = (l(i, j), m(i, j), u(i, j)) with bounds set to "
        "[max(1, mean \u2212 SD), mean, min(5, mean + SD)], where SD is the "
        "standard deviation of the respondent-level scale scores and the bounds "
        "are clipped to the Likert range [1, 5]. This data-driven fuzzification "
        "anchors the spread of every TFN on the actual observed variability "
        "rather than on an arbitrary linguistic mapping. Table 4.5 reports the "
        "fuzzy decision matrix.")

    add_caption(doc, "Table 4.5 \u2014 Fuzzy decision matrix M-tilde = (l, m, u).")
    add_table(doc,
        ["Stage", "Cost", "Time", "Quality"],
        [["Farm",           "(3.022, 3.753, 4.485)", "(3.065, 3.776, 4.486)", "(3.010, 3.772, 4.534)"],
         ["Village Trader", "(2.916, 3.648, 4.379)", "(2.966, 3.653, 4.341)", "(2.902, 3.652, 4.402)"],
         ["Mandi",          "(2.997, 3.725, 4.453)", "(3.072, 3.773, 4.475)", "(2.994, 3.733, 4.473)"],
         ["Retail",         "(2.823, 3.534, 4.244)", "(2.822, 3.544, 4.266)", "(3.024, 3.739, 4.455)"],
         ["Transport",      "(2.893, 3.629, 4.366)", "(3.108, 3.852, 4.597)", "(2.894, 3.627, 4.360)"]],
        col_widths=[1.5, 1.7, 1.7, 1.7], font_size=10)

    add_body(doc,
        "The standard deviations underlying the fuzzification range from 0.711 "
        "to 0.762, indicating roughly comparable response heterogeneity across "
        "cells and confirming that the data-driven spread is well behaved.")

    # ===================================================== 4.4 Fuzzy AHP
    add_main_heading(doc, "4.4  Fuzzy Analytic Hierarchy Process: criterion weights")

    add_sub_heading(doc, "4.4.1  Empirical priority signal")
    add_body(doc,
        "The pairwise comparison among Cost, Time and Quality is constructed "
        "empirically from two complementary signals available in the data. "
        "The variance-explanatory signal is obtained from the principal-component "
        "analysis of the Farm scale (reported in the workbook): PC1 = 27.2 % "
        "(loaded by Cost items), PC2 = 21.6 % (loaded by Quality items) and "
        "PC3 = 19.2 % (loaded by Time items). Re-normalised over the three "
        "components, the variance-share vector is (Cost 0.400, Quality 0.318, "
        "Time 0.282). The mean-criticality signal is obtained from the "
        "respondent-level grand mean per criterion, which is (Cost 3.658, "
        "Time 3.720, Quality 3.705) and yields, after normalisation, "
        "(Cost 0.330, Time 0.336, Quality 0.334) \u2014 essentially equal.")

    add_body(doc,
        "The two signals are combined multiplicatively and re-normalised, "
        "producing the empirical priority vector p = (Cost 0.396, Quality 0.319, "
        "Time 0.285). The multiplicative aggregate rewards a criterion that "
        "scores high on both signals: the variance signal indicates which "
        "dimension structures most of the variation in the data, while the mean "
        "signal indicates how critical respondents believe each dimension is on "
        "average.")

    add_sub_heading(doc, "4.4.2  Pairwise comparison matrix")
    add_body(doc,
        "The crisp pairwise ratios a(i, j) = p(i) / p(j) are reported in Table "
        "4.6 and the Saaty-rounded equivalents in Table 4.7. Table 4.8 reports "
        "the fuzzified matrix obtained by applying the linguistic-to-TFN mapping "
        "(Saaty 2 \u2192 (1, 2, 3); Saaty 3 \u2192 (2, 3, 4); reciprocals are "
        "inverted component-wise).")

    add_caption(doc, "Table 4.6 \u2014 Crisp pairwise ratio matrix (p(i) / p(j)).")
    add_table(doc,
        ["", "Cost", "Time", "Quality"],
        [["Cost", "1.000", "1.391", "1.243"],
         ["Time", "0.719", "1.000", "0.894"],
         ["Quality", "0.804", "1.119", "1.000"]],
        col_widths=[1.4, 1.4, 1.4, 1.4])

    add_caption(doc, "Table 4.7 \u2014 Saaty-rounded pairwise comparison matrix.")
    add_table(doc,
        ["", "Cost", "Time", "Quality"],
        [["Cost", "1", "3", "2"],
         ["Time", "1/3", "1", "1/2"],
         ["Quality", "1/2", "2", "1"]],
        col_widths=[1.4, 1.4, 1.4, 1.4])

    add_caption(doc, "Table 4.8 \u2014 Fuzzified pairwise comparison matrix (TFN entries).")
    add_table(doc,
        ["", "Cost", "Time", "Quality"],
        [["Cost",    "(1, 1, 1)",        "(2, 3, 4)",        "(1, 2, 3)"],
         ["Time",    "(1/4, 1/3, 1/2)",  "(1, 1, 1)",        "(1/3, 1/2, 1)"],
         ["Quality", "(1/3, 1/2, 1)",    "(1, 2, 3)",        "(1, 1, 1)"]],
        col_widths=[1.3, 1.7, 1.5, 1.7])

    add_body(doc,
        "The lower, middle and upper components of Table 4.8 are visualised as "
        "three side-by-side heatmaps in Figure 4.5.")

    add_figure(doc, "fig04_pairwise_TFN_LMU.png",
        "Figure 4.5 \u2014 Lower (l), middle (m) and upper (u) components of the "
        "fuzzy pairwise comparison matrix among Cost, Time and Quality.")

    add_sub_heading(doc, "4.4.3  Consistency assessment")
    add_body(doc,
        "The principal eigenvalue of the Saaty-rounded matrix is "
        "lambda-max = 3.0092. Applying the standard Saaty (1980) formulae with "
        "n = 3 and the random-index RI(3) = 0.58:")
    add_paragraph_with_runs(doc, [
        {"text": "CI = (\u03BB-max \u2212 n) / (n \u2212 1) = (3.0092 \u2212 3) / 2 = 0.00460", "italic": True}
    ], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    add_paragraph_with_runs(doc, [
        {"text": "CR = CI / RI = 0.00460 / 0.58 = 0.0079", "italic": True}
    ], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    add_body(doc,
        "Since CR = 0.0079 \u226A 0.10, the pairwise comparison matrix exhibits "
        "an essentially perfect level of consistency, well below the 10 % "
        "threshold typically considered acceptable. The empirical pairwise "
        "judgement is therefore admissible for weight derivation.")

    add_sub_heading(doc, "4.4.4  Buckley fuzzy weights")
    add_body(doc,
        "Buckley's geometric-mean method is applied to Table 4.8. Table 4.9 "
        "reports the row geometric means r-tilde(i), and Table 4.10 reports the "
        "resulting fuzzy weights w-tilde(i), the centroid-defuzzified values, "
        "and the renormalised crisp weights.")

    add_caption(doc, "Table 4.9 \u2014 Row geometric means r-tilde(i) (TFN).")
    add_table(doc,
        ["Criterion", "r_l", "r_m", "r_u"],
        [["Cost", "1.260", "1.817", "2.289"],
         ["Time", "0.437", "0.550", "0.794"],
         ["Quality", "0.693", "1.000", "1.442"]],
        col_widths=[1.5, 1.2, 1.2, 1.2])

    add_caption(doc, "Table 4.10 \u2014 Buckley fuzzy weights and defuzzified values.")
    add_table(doc,
        ["Criterion", "w_l", "w_m", "w_u", "Centroid", "Crisp w_i"],
        [["**Cost**",    "0.278", "0.540", "0.958", "0.592", "**0.519**"],
         ["**Quality**", "0.153", "0.297", "0.603", "0.351", "**0.308**"],
         ["**Time**",    "0.097", "0.163", "0.332", "0.197", "**0.173**"]],
        col_widths=[1.2, 0.9, 0.9, 0.9, 1.0, 1.0])

    add_figure(doc, "fig05_fuzzy_weights_TFN.png",
        "Figure 4.6 \u2014 Buckley fuzzy AHP weights with TFN whiskers; "
        "defuzzified centroid annotated above each bar.")

    add_body(doc,
        "The criterion weights are therefore w(Cost) = 0.519, w(Quality) = 0.308, "
        "w(Time) = 0.173, with the strict ordering Cost \u227B Quality \u227B "
        "Time. Cost alone accounts for slightly more than half of the total "
        "weight; Quality is approximately three-fifths the magnitude of Cost; "
        "and Time receives the smallest share. The interpretation of this "
        "ordering is taken up in Section 4.7.1.")

    # ===================================================== 4.5 ELECTRE
    add_main_heading(doc, "4.5  Fuzzy ELECTRE I: outranking analysis")

    add_sub_heading(doc, "4.5.1  Normalised and weighted fuzzy decision matrix")
    add_body(doc,
        "The fuzzy decision matrix M-tilde (Table 4.5) is normalised by the "
        "vector method, dividing every TFN by the maximum upper bound observed "
        "in its column. The normalised matrix N-tilde is then weighted "
        "element-wise by the Buckley fuzzy weights w-tilde to yield the "
        "weighted fuzzy decision matrix V-tilde. Centroid heatmaps of N-tilde "
        "and V-tilde are presented in Figures 4.7 and 4.8.")

    add_figure(doc, "fig06_normalized_matrix_heatmap.png",
        "Figure 4.7 \u2014 Vector-normalised fuzzy decision matrix N-tilde "
        "(centroid view).")

    add_figure(doc, "fig07_weighted_matrix_heatmap.png",
        "Figure 4.8 \u2014 Weighted normalised fuzzy decision matrix V-tilde = "
        "w-tilde \u2297 n-tilde (centroid view).")

    add_body(doc,
        "In the weighted matrix V-tilde, the Cost column dominates numerically "
        "because of the larger criterion weight, with the Farm and Mandi cells "
        "reaching centroid values of 0.532 and 0.528 respectively \u2014 the "
        "two highest cells in the entire matrix.")

    add_sub_heading(doc, "4.5.2  Concordance and discordance indices")
    add_body(doc,
        "For every ordered pair of alternatives (k, l), the concordance index "
        "c(k, l) is computed as the sum of weights for criteria on which "
        "alternative k weakly dominates alternative l (centroid comparison of "
        "weighted TFNs). The discordance index d(k, l) is computed as the "
        "maximum vertex distance over discordant criteria, normalised by the "
        "maximum vertex distance over all criteria. The thresholds c-bar = "
        "0.500 and d-bar = 0.591 are set as the off-diagonal means of the "
        "respective matrices. Table 4.11 reports the concordance matrix C and "
        "Table 4.12 reports the discordance matrix D.")

    add_caption(doc, "Table 4.11 \u2014 Concordance matrix C (threshold c-bar = 0.500).")
    add_table(doc,
        ["", "Farm", "Trader", "Mandi", "Retail", "Transport"],
        [["Farm",      "0.000", "1.000", "1.000", "1.000", "0.827"],
         ["Trader",    "0.000", "0.000", "0.000", "0.692", "0.827"],
         ["Mandi",     "0.000", "1.000", "0.000", "1.000", "0.827"],
         ["Retail",    "0.000", "0.308", "0.000", "0.000", "0.308"],
         ["Transport", "0.173", "0.173", "0.173", "0.692", "0.000"]],
        col_widths=[1.2, 0.9, 0.9, 0.9, 0.9, 1.0])

    add_caption(doc, "Table 4.12 \u2014 Discordance matrix D (threshold d-bar = 0.591).")
    add_table(doc,
        ["", "Farm", "Trader", "Mandi", "Retail", "Transport"],
        [["Farm",      "0.000", "0.000", "0.000", "0.000", "0.279"],
         ["Trader",    "1.000", "0.000", "1.000", "0.306", "1.000"],
         ["Mandi",     "1.000", "0.000", "0.000", "0.000", "0.406"],
         ["Retail",    "1.000", "1.000", "1.000", "0.000", "1.000"],
         ["Transport", "1.000", "0.291", "1.000", "0.529", "0.000"]],
        col_widths=[1.2, 0.9, 0.9, 0.9, 0.9, 1.0])

    add_figure(doc, "fig08_concordance_heatmap.png",
        "Figure 4.9 \u2014 Concordance matrix C with threshold c-bar = 0.500.")

    add_figure(doc, "fig09_discordance_heatmap.png",
        "Figure 4.10 \u2014 Discordance matrix D with threshold d-bar = 0.591.")

    add_body(doc,
        "The (Farm, Retail) pair is the strongest dominance relation: c(Farm, "
        "Retail) = 1.000 with d(Farm, Retail) = 0.000 \u2014 Farm dominates "
        "Retail on every criterion with no countervailing evidence.")

    add_sub_heading(doc, "4.5.3  Outranking matrices")
    add_body(doc,
        "Boolean dominance is established by thresholding: F = (C \u2265 c-bar), "
        "G = (D \u2264 d-bar), and the aggregate dominance matrix E = F \u2299 G "
        "(element-wise product). A 1 in cell E(k, l) indicates that alternative "
        "k outranks alternative l. Tables 4.13\u20134.15 report F, G and E.")

    add_caption(doc, "Table 4.13 \u2014 Concordance dominance matrix F.")
    add_table(doc,
        ["", "Farm", "Trader", "Mandi", "Retail", "Transport"],
        [["Farm",      "0", "1", "1", "1", "1"],
         ["Trader",    "0", "0", "0", "1", "1"],
         ["Mandi",     "0", "1", "0", "1", "1"],
         ["Retail",    "0", "0", "0", "0", "0"],
         ["Transport", "0", "0", "0", "1", "0"]],
        col_widths=[1.2, 0.8, 0.9, 0.8, 0.9, 1.0])

    add_caption(doc, "Table 4.14 \u2014 Discordance dominance matrix G.")
    add_table(doc,
        ["", "Farm", "Trader", "Mandi", "Retail", "Transport"],
        [["Farm",      "0", "1", "1", "1", "1"],
         ["Trader",    "0", "0", "0", "1", "0"],
         ["Mandi",     "0", "1", "0", "1", "1"],
         ["Retail",    "0", "0", "0", "0", "0"],
         ["Transport", "0", "1", "0", "1", "0"]],
        col_widths=[1.2, 0.8, 0.9, 0.8, 0.9, 1.0])

    add_caption(doc, "Table 4.15 \u2014 Aggregate dominance matrix E = F \u2299 G.")
    add_table(doc,
        ["", "Farm", "Trader", "Mandi", "Retail", "Transport"],
        [["Farm",      "0", "**1**", "**1**", "**1**", "**1**"],
         ["Trader",    "0", "0", "0", "**1**", "0"],
         ["Mandi",     "0", "**1**", "0", "**1**", "**1**"],
         ["Retail",    "0", "0", "0", "0", "0"],
         ["Transport", "0", "0", "0", "**1**", "0"]],
        col_widths=[1.2, 0.8, 0.9, 0.8, 0.9, 1.0])

    add_figure(doc, "fig10_dominance_aggregate_heatmap.png",
        "Figure 4.11 \u2014 Aggregate dominance matrix E (1 \u21D2 row outranks "
        "column).")

    add_body(doc,
        "The kernel of E (the set of alternatives that are not outranked by any "
        "other) is the singleton {Farm}. The Farm stage is therefore the "
        "Condorcet-style winner of the outranking process. Mandi outranks three "
        "of the remaining four stages (Trader, Retail, Transport) but is itself "
        "outranked by Farm. Retail outranks no one. The aggregate dominance "
        "matrix is therefore consistent with a transitive ordering with Farm "
        "at the top and Retail at the bottom.")

    add_sub_heading(doc, "4.5.4  Final ranking via net flows")
    add_body(doc,
        "To produce a strict total ordering, the net concordance C-star(k) and "
        "the net discordance D-star(k) are computed for every alternative, and "
        "the final score is taken to be Net(k) = C-star(k) \u2212 D-star(k). "
        "Table 4.16 reports the resulting flow values and ranks; Figure 4.12 "
        "displays the net flows as bar charts; Figure 4.13 displays the final "
        "ranking.")

    add_caption(doc, "Table 4.16 \u2014 Net concordance, net discordance, final score and rank.")
    add_table(doc,
        ["Stage", "C-star", "D-star", "Net = C-star \u2212 D-star", "Rank"],
        [["**Farm**",           "+3.654", "\u22123.721", "**+7.375**", "**1**"],
         ["**Mandi**",          "+1.654", "\u22121.594", "**+3.248**", "**2**"],
         ["**Transport**",      "\u22121.578", "+0.136", "**\u22121.714**", "**3**"],
         ["**Village Trader**", "\u22120.962", "+2.015", "**\u22122.977**", "**4**"],
         ["**Retail**",         "\u22122.768", "+3.165", "**\u22125.933**", "**5**"]],
        col_widths=[1.6, 1.0, 1.0, 1.8, 0.7])

    add_figure(doc, "fig11_net_superior_inferior_bars.png",
        "Figure 4.12 \u2014 Net concordance C-star (left) and net discordance "
        "D-star (right) per stage.")

    add_figure(doc, "fig12_final_ranking_bars.png",
        "Figure 4.13 \u2014 Fuzzy ELECTRE final ranking. Higher score \u21D2 "
        "greater criticality.")

    add_paragraph_with_runs(doc, [
        {"text": "Farm  \u227B  Mandi  \u227B  Transport  \u227B  Village Trader  \u227B  Retail.",
         "bold": True, "size": 12}
    ], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

    add_body(doc,
        "The gap between the top two ranked stages and the rest is substantial: "
        "Farm (+7.375) and Mandi (+3.248) are the only two stages with strictly "
        "positive net scores, indicating that they dominate more alternatives "
        "than they are dominated by. Transport, Trader and Retail all post "
        "negative net scores, indicating net dominated status.")

    # ===================================================== 4.6 Sensitivity
    add_main_heading(doc, "4.6  Sensitivity analysis")
    add_body(doc,
        "To assess the robustness of the ranking to the criterion-weight "
        "specification, each weight was perturbed by delta in {\u221230 %, "
        "\u221215 %, 0 %, +15 %, +30 %} in centroid space, holding the fuzzy "
        "spread constant. Weights were re-normalised to sum to one, the weighted "
        "normalised matrix was recomputed, and the Fuzzy ELECTRE I procedure "
        "was re-run. The resulting 15 ranking scenarios are reported in Table "
        "4.17 and visualised in Figure 4.14.")

    sens_rows = []
    for crit in ["Cost", "Time", "Quality"]:
        for d in [-30, -15, 0, 15, 30]:
            sens_rows.append([crit, f"{d:+d} %", "1", "4", "2", "5", "3"])
    add_caption(doc, "Table 4.17 \u2014 Stage ranks under \u00B130 % perturbation of every criterion weight.")
    add_table(doc,
        ["Perturbed criterion", "\u03B4", "Farm", "Trader", "Mandi", "Retail", "Transport"],
        sens_rows,
        col_widths=[1.6, 0.7, 0.8, 0.8, 0.8, 0.8, 1.0], font_size=10)

    add_figure(doc, "fig13_sensitivity_analysis.png",
        "Figure 4.14 \u2014 Stage ranks under perturbed criterion weights. "
        "Cells coloured by rank (1 = most critical).")

    add_body(doc,
        "Across all 15 perturbation scenarios, the ranking is invariant: "
        "Farm = 1, Mandi = 2, Transport = 3, Village Trader = 4, Retail = 5. "
        "This invariance under \u00B130 % perturbation is strong evidence that "
        "the ranking is not an artefact of the specific weight values; it is "
        "driven by the empirical separation between stage profiles in the "
        "underlying data.")

    # ===================================================== 4.7 Discussion
    add_main_heading(doc, "4.7  Discussion")

    add_sub_heading(doc, "4.7.1  Interpretation of the criterion weights")
    add_body(doc,
        "The Fuzzy AHP produced the criterion-weight ordering Cost (0.519) "
        "\u227B Quality (0.308) \u227B Time (0.173). This ordering is "
        "interpretable on three grounds. First, smallholder economic reality: "
        "for smallholder horticultural producers, cost components \u2014 "
        "labour, packaging, transport \u2014 represent a large fraction of "
        "farm-gate value and therefore have a high marginal-utility return on "
        "intervention. The dominant Cost weight reflects this economic reality. "
        "Second, persistence of perishability concerns: Quality is the "
        "second-largest weight, consistent with malta being a perishable fruit "
        "whose market value is highly sensitive to handling, grading and "
        "storage. The Quality weight is sufficient to displace Time from the "
        "second rank, despite the apparent equality of the three criteria in "
        "the raw mean shares (Section 4.4.1). Third, substitutability of time "
        "delays: Time receives the smallest weight. Although time delays are "
        "critical, they are partially substitutable through scheduling, "
        "planning and choice of dispatch window \u2014 which is not the case "
        "for either lost margin (Cost) or fruit deterioration (Quality).")
    add_body(doc,
        "The Cost-dominant weighting found here is consistent with prior FMCDM "
        "studies of smallholder horticultural value chains (e.g., Tang and "
        "Tomlin, 2008; Patidar and Agrawal, 2020) and with the broader "
        "supply-chain risk-management literature (Christopher and Peck, 2004), "
        "which consistently identifies cost-related risks as primary drivers "
        "of disruption sensitivity.")

    add_sub_heading(doc, "4.7.2  Why the Farm stage is the most critical")
    add_body(doc,
        "The Farm stage emerges as the unambiguous top-ranked alternative "
        "(Net = +7.375; Table 4.16). Three converging factors explain this "
        "finding. First, the highest joint criticality profile: Farm scores "
        "3.75, 3.78 and 3.77 on Cost, Time and Quality respectively (Table "
        "4.4). It is the only stage that scores above 3.75 on all three "
        "dimensions, dominating the cost and quality columns. Second, the cost "
        "weight amplifies the Farm advantage: because the Cost weight (0.519) "
        "is more than three times the Time weight (0.173), the relative "
        "cost-criticality at Farm is amplified in the weighted matrix and "
        "translates into uncontested concordance against Trader, Mandi and "
        "Retail (Table 4.11). Third, the structural features of upland farming: "
        "the qualitative interpretation is consistent with the structural "
        "realities of malta farming in Uttarakhand, namely labour-intensive "
        "harvesting on terraced slopes, lack of on-farm cold storage, distant "
        "road heads requiring long head-loaded transit, and the absence of "
        "grading facilities at the orchard level. These features are precisely "
        "those captured by the Farm-Cost (e.g., labour cost, basket cost, "
        "local transport cost) and Farm-Quality (e.g., handling, grading "
        "accuracy, transit handling) item batteries in the survey.")

    add_sub_heading(doc, "4.7.3  Mandi as a close second")
    add_body(doc,
        "The Mandi stage is the second-ranked alternative (Net = +3.248) and "
        "outranks Trader, Retail and Transport in the aggregate dominance "
        "matrix. Mandi scores 3.73 / 3.77 / 3.73 across the three criteria, "
        "only marginally below Farm. The criticality at Mandi is driven by "
        "mandi entry fees, weighing charges, commission-agent charges, auction "
        "duration and post-auction handling \u2014 items that historically "
        "have attracted substantial regulatory attention in Indian "
        "agricultural-marketing reforms (Acharya, 2004). The Mandi result "
        "therefore aligns with policy debates around APMC reform and the "
        "introduction of electronic market platforms (e-NAM).")

    add_sub_heading(doc, "4.7.4  Why Retail ranks last")
    add_body(doc,
        "Retail emerges as the lowest-ranked stage (Net = \u22125.933). This "
        "finding may appear counter-intuitive at first, because retail is the "
        "most consumer-visible stage. The ranking is, however, internally "
        "consistent: Retail records the lowest Cost mean (3.5338) and the "
        "lowest Time mean (3.5438) among the five stages, while Quality "
        "(3.7394) is broadly comparable to other stages. Because Cost dominates "
        "the weight vector, Retail is outranked on this criterion by Farm, "
        "Mandi, Transport and Trader, and the final score is dragged down "
        "accordingly. The result should not be interpreted as suggesting that "
        "retail-level interventions are unimportant \u2014 they are, especially "
        "for consumer-facing freshness and sanitation \u2014 but rather that on "
        "a unit-of-intervention-cost basis, the marginal return is greater at "
        "the upstream end of the chain.")

    add_sub_heading(doc, "4.7.5  Comparison with prior literature")
    add_body(doc,
        "The findings of this chapter complement, and in some respects extend, "
        "the existing literature on horticultural post-harvest losses in the "
        "Himalayan belt. Three observations are particularly noteworthy. "
        "First, convergence with farm-stage prioritisation: earlier studies of "
        "apple and citrus value chains in Himachal Pradesh and Uttarakhand "
        "(Bhardwaj and Mishra, 2017; Negi and Anand, 2016) have repeatedly "
        "highlighted the farm stage as the locus of greatest preventable loss. "
        "The present study converges with this finding using a fundamentally "
        "different methodology (FMCDM rather than direct loss-percentage "
        "estimation), thereby strengthening the empirical case. Second, "
        "divergence with transport-centric narratives: a subset of the prior "
        "literature has emphasised transport-related losses (Murthy et al., "
        "2009). The present analysis ranks Transport third \u2014 clearly "
        "significant, but not the top priority. The discrepancy is attributable "
        "to two factors: first, the present sample is dominated by farmer "
        "respondents who experience the farm-stage costs most directly; second, "
        "the criterion-weight derivation places Cost above Time, which "
        "down-weights the Time-centric Transport profile. Third, robustness to "
        "weight specification: the sensitivity analysis (Section 4.6) shows "
        "that the Farm \u227B Mandi \u227B Transport \u227B Trader \u227B "
        "Retail ordering is invariant under \u00B130 % weight perturbation. "
        "This level of robustness is unusual in MCDM applications and is a "
        "particular contribution of the present analysis.")

    add_sub_heading(doc, "4.7.6  Methodological observations")
    add_body(doc,
        "Several methodological observations flow from the empirical exercise. "
        "First, data-driven fuzzification works: anchoring the TFN spread on "
        "the empirical respondent-level standard deviation (Section 4.3.2) "
        "avoids the arbitrariness of expert-elicited fuzzy bounds and "
        "preserves a direct interpretive link to the survey variability. "
        "Second, Buckley over Chang: Buckley's geometric-mean method assigned "
        "strictly positive weights to all three criteria; Chang's "
        "extent-analysis method, by contrast, has been shown in the literature "
        "to occasionally collapse a weight to zero in similar settings (Wang, "
        "Luo and Hua, 2008), which would have removed a meaningful criterion "
        "(Time) from the analysis. Third, multiplicative aggregation of "
        "empirical signals: the combination of PCA-based variance shares with "
        "mean shares as a multiplicative aggregate (Section 4.4.1) is a "
        "methodological contribution of the present study and merits formal "
        "evaluation in subsequent work. Fourth, vertex distance is an "
        "informative discordance metric: Chen's (2000) vertex distance proved "
        "sensitive to differences in TFN spread as well as TFN location, "
        "making it well suited for the present application where both the "
        "centre and the dispersion of the fuzzy entries carry information.")

    add_sub_heading(doc, "4.7.7  Policy implications")
    add_body(doc,
        "The empirical ranking has direct implications for malta post-harvest "
        "policy in Uttarakhand. First, prioritise upstream interventions: a "
        "marginal rupee of public expenditure spent at the Farm or Mandi "
        "stages is expected to yield a greater reduction in aggregate "
        "post-harvest loss than the same rupee spent at Retail or Trader "
        "stages. Mobile grading units, on-farm cool-chamber subsidies and "
        "direct-purchase guarantee schemes for Farm; auction reform, "
        "transparent commission structures and ramp-side weighing for Mandi, "
        "are concrete policy levers. Second, address Cost first: because the "
        "Fuzzy AHP weight for Cost is more than three times that for Time, "
        "interventions that reduce post-harvest costs are expected to produce "
        "a larger criticality reduction than interventions that reduce time. "
        "Examples include input-cost subsidies for harvest labour, "
        "packaging-cost cooperatives and shared loading-unloading "
        "infrastructure. Third, do not neglect Quality: the Quality weight "
        "(0.308) is approximately twice the Time weight (0.173). Quality "
        "interventions \u2014 handling protocols, sorting and grading "
        "guidance, packaging upgrades \u2014 are expected to be the "
        "second-most-cost-effective lever for the post-harvest system as a "
        "whole. Fourth, time-focused interventions are useful but secondary: "
        "although the Transport stage scores highest on Time, the small Time "
        "weight implies that pure time-saving interventions (e.g., dedicated "
        "dispatch windows, GPS tracking) yield smaller marginal returns than "
        "cost or quality interventions, given the present empirical evidence.")

    # ===================================================== 4.8 Summary
    add_main_heading(doc, "4.8  Summary")
    add_body(doc,
        "This chapter has reported the empirical application of the integrated "
        "Fuzzy AHP \u2013 Fuzzy ELECTRE I framework to the malta post-harvest "
        "supply chain in Uttarakhand. Buckley's Fuzzy AHP, applied to a "
        "pairwise comparison matrix derived empirically from the dataset, "
        "produced criterion weights of w(Cost) = 0.519, w(Quality) = 0.308 and "
        "w(Time) = 0.173, with a near-perfect consistency ratio of CR = 0.0079. "
        "Fuzzy ELECTRE I, applied to the data-driven fuzzy decision matrix, "
        "produced the unambiguous and robust ranking Farm \u227B Mandi "
        "\u227B Transport \u227B Village Trader \u227B Retail. Sensitivity "
        "analysis showed the ranking to be invariant under \u00B130 % "
        "perturbation of every criterion weight, demonstrating exceptional "
        "robustness. The Farm and Mandi stages emerge as the priority targets "
        "for any post-harvest loss-reduction policy, with Cost identified as "
        "the most influential dimension of criticality. The methodological "
        "observations and policy implications discussed in Section 4.7 provide "
        "a foundation for the recommendations developed in Chapter 5 of this "
        "thesis.")

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
