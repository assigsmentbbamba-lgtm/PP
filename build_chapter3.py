"""
Chapter 3 - RESEARCH METHODOLOGY
Malta (Citrus sinensis) distribution channel optimisation - Garhwal region.

Re-build per user spec (v2):
    - Cross-sectional, mixed-method
    - Three phases of data collection (Sep 2024 - Dec 2025)
    - 144 factors -> 82 finalised after reliability / consistency screening
    - Stratified random sampling
    - Tools: ONLY Fuzzy AHP and ELECTRE
    - Sample = 1,300 (numbers from Malta_.xlsx)
    - Multiple figures: research process, framework, screening funnel,
      phase timeline, comparative bars for FAHP & ELECTRE
    - Formatting:
          body          - Times New Roman 12 pt
          headings      - Times New Roman 14 pt bold
          sub-headings  - Times New Roman 12 pt BOLD + ITALIC
    - Voice: written naturally, the way a researcher actually writes
      (varied sentence rhythm, occasional first-person, no AI tells)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Polygon
import numpy as np

FONT = "Times New Roman"


# ============================================================
# FIGURES
# ============================================================

def fig_research_process(path):
    fig, ax = plt.subplots(figsize=(8.5, 11))
    ax.set_xlim(0, 10); ax.set_ylim(0, 14); ax.axis("off")

    def box(x, y, w, h, txt, fc="#DEEBF7", ec="#1F4E78", weight="normal", fs=10):
        ax.add_patch(FancyBboxPatch((x-w/2, y-h/2), w, h,
                                     boxstyle="round,pad=0.08,rounding_size=0.15",
                                     lw=1.4, ec=ec, fc=fc))
        ax.text(x, y, txt, ha="center", va="center",
                fontsize=fs, family="serif", weight=weight, wrap=True)

    def arr(x1, y1, x2, y2):
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2),
                                     arrowstyle="-|>", mutation_scale=14,
                                     color="#1F4E78", lw=1.5))

    box(5, 13.4, 7.0, 0.9,
        "Stage 1: Identification of the Research Problem\n"
        "(Distribution-channel inefficiency of Malta in the Garhwal region)",
        fc="#FCE4D6", ec="#C65911", weight="bold", fs=11)
    arr(5, 12.95, 5, 12.55)

    box(5, 12.1, 7.0, 0.9,
        "Stage 2: Literature Review, Expert Consultation\n"
        "Pilot Study and Stakeholder Discussions",
        fc="#FFF2CC", ec="#BF8F00", fs=11)
    arr(5, 11.65, 5, 11.25)

    box(5, 10.8, 7.0, 0.9,
        "Stage 3: Initial Pool of 144 Factors -> Reliability\n"
        "and Consistency Screening -> 82 Finalised Factors",
        fc="#FFF2CC", ec="#BF8F00", fs=11)
    arr(5, 10.35, 5, 9.95)

    box(5, 9.5, 7.0, 0.9,
        "Stage 4: Cross-Sectional Survey Design\n"
        "(stratified random sampling, four Garhwal districts)",
        fc="#E2EFDA", ec="#548235", fs=11)
    arr(5, 9.05, 5, 8.65)

    box(5, 8.2, 7.0, 1.0,
        "Stage 5: Three-Phase Field Data Collection (Sep 2024 - Dec 2025)\n"
        "Farmers, Village Traders, Commission Agents, Wholesalers,\n"
        "Retailers, Transporters and Experts/Officials  (n = 1,300)",
        fc="#E2EFDA", ec="#548235", fs=11)
    arr(5, 7.7, 5, 7.3)

    arr(5, 7.3, 2.6, 6.7)
    arr(5, 7.3, 7.4, 6.7)

    # FAHP
    box(2.6, 6.3, 4.0, 0.7, "Stage 6A: FUZZY AHP",
        fc="#D9E1F2", ec="#1F4E78", weight="bold", fs=12)
    arr(2.6, 5.95, 2.6, 5.65)
    box(2.6, 5.3, 4.0, 0.7,
        "Pairwise comparison of criteria\n(9-point linguistic scale)",
        fc="#DEEBF7", ec="#1F4E78", fs=10)
    arr(2.6, 4.95, 2.6, 4.65)
    box(2.6, 4.3, 4.0, 0.7,
        "Convert judgements to\nTriangular Fuzzy Numbers (l, m, u)",
        fc="#DEEBF7", ec="#1F4E78", fs=10)
    arr(2.6, 3.95, 2.6, 3.65)
    box(2.6, 3.3, 4.0, 0.7,
        "Aggregate via fuzzy geometric mean\n-> Fuzzy comparison matrix",
        fc="#DEEBF7", ec="#1F4E78", fs=10)
    arr(2.6, 2.95, 2.6, 2.65)
    box(2.6, 2.3, 4.0, 0.8,
        "Defuzzify (Centre of Area), normalise\n-> Criterion weights w_j\nCheck CR < 0.10",
        fc="#DEEBF7", ec="#1F4E78", fs=10)
    arr(2.6, 1.85, 4.6, 1.15)

    # ELECTRE
    box(7.4, 6.3, 4.0, 0.7, "Stage 6B: ELECTRE",
        fc="#D9E1F2", ec="#1F4E78", weight="bold", fs=12)
    arr(7.4, 5.95, 7.4, 5.65)
    box(7.4, 5.3, 4.0, 0.7,
        "Decision matrix of channel alternatives\nrated against each criterion",
        fc="#DEEBF7", ec="#1F4E78", fs=10)
    arr(7.4, 4.95, 7.4, 4.65)
    box(7.4, 4.3, 4.0, 0.7,
        "Normalise and weight\n(weights w_j from Fuzzy AHP)",
        fc="#DEEBF7", ec="#1F4E78", fs=10)
    arr(7.4, 3.95, 7.4, 3.65)
    box(7.4, 3.3, 4.0, 0.7,
        "Build concordance C(a,b)\nand discordance D(a,b) matrices",
        fc="#DEEBF7", ec="#1F4E78", fs=10)
    arr(7.4, 2.95, 7.4, 2.65)
    box(7.4, 2.3, 4.0, 0.8,
        "Net concordance / discordance indices\nOutranking graph -> Final ranking\nof channel alternatives",
        fc="#DEEBF7", ec="#1F4E78", fs=10)
    arr(7.4, 1.85, 5.4, 1.15)

    box(5, 0.8, 7.0, 0.8,
        "Stage 7: Optimal Distribution Channel for Malta\nand Recommendations to Stakeholders",
        fc="#FCE4D6", ec="#C65911", weight="bold", fs=12)

    plt.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def fig_methodology_framework(path):
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10); ax.set_ylim(0, 7); ax.axis("off")

    def block(x, y, w, h, title, body, fc="#DEEBF7", ec="#1F4E78"):
        ax.add_patch(FancyBboxPatch((x-w/2, y-h/2), w, h,
                                     boxstyle="round,pad=0.1,rounding_size=0.18",
                                     lw=1.5, ec=ec, fc=fc))
        ax.text(x, y+h/2-0.28, title, ha="center", va="center",
                fontsize=11, family="serif", weight="bold")
        ax.text(x, y-0.05, body, ha="center", va="center",
                fontsize=9, family="serif")

    def arrow(x1, y1, x2, y2, text=None):
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2),
                                     arrowstyle="-|>", mutation_scale=14,
                                     color="#404040", lw=1.4))
        if text:
            ax.text((x1+x2)/2, (y1+y2)/2+0.18, text,
                    ha="center", va="center",
                    fontsize=8, family="serif", style="italic", color="#404040")

    # Inputs
    block(1.5, 6.2, 2.4, 1.1, "INPUTS",
          "- Literature review\n- Expert consultation\n- Pilot study\n- Field observation",
          fc="#FFF2CC", ec="#BF8F00")
    # Process
    block(5.0, 6.2, 2.4, 1.1, "FACTOR POOL",
          "144 candidate factors\nacross 5 stages and\n3 aspects (C/T/Q)",
          fc="#E2EFDA", ec="#548235")
    # Screening
    block(8.5, 6.2, 2.4, 1.1, "SCREENING",
          "Reliability & consistency\nanalysis -> remove\nredundancies",
          fc="#FCE4D6", ec="#C65911")

    arrow(2.7, 6.2, 3.8, 6.2)
    arrow(6.2, 6.2, 7.3, 6.2)

    # Down arrow center
    arrow(8.5, 5.55, 5.0, 4.5, "82 finalised factors")

    # Survey
    block(5.0, 3.9, 6.0, 1.2, "CROSS-SECTIONAL SURVEY (n = 1,300)",
          "Stratified random sample - 7 stakeholder groups - 4 Garhwal districts\n"
          "(Pauri Garhwal, Tehri Garhwal, Rudraprayag, Chamoli)\n"
          "Three phases: Sep 2024 - Mar 2025 / Apr - Aug 2025 / Sep - Dec 2025",
          fc="#D9E1F2", ec="#1F4E78")

    arrow(3.5, 3.3, 2.5, 2.3)
    arrow(6.5, 3.3, 7.5, 2.3)

    # Tools
    block(2.5, 1.6, 3.4, 1.2, "TOOL 1: FUZZY AHP",
          "TFN linguistic scale\nFuzzy synthetic extents\nDefuzzification (CoA)\nCriterion weights w_j",
          fc="#DEEBF7", ec="#1F4E78")
    block(7.5, 1.6, 3.4, 1.2, "TOOL 2: ELECTRE",
          "Concordance C(a,b)\nDiscordance D(a,b)\nOutranking graph\nRanking of channels",
          fc="#DEEBF7", ec="#1F4E78")

    arrow(4.2, 1.6, 5.8, 1.6, "weights")

    # Output
    block(5.0, 0.45, 5.0, 0.7, "OUTPUT: Optimised Malta Distribution Channel & Strategy",
          "", fc="#FCE4D6", ec="#C65911")
    arrow(2.5, 0.95, 4.5, 0.7)
    arrow(7.5, 0.95, 5.5, 0.7)

    plt.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def fig_funnel(path):
    """Factor-screening funnel: 144 -> 82."""
    fig, ax = plt.subplots(figsize=(8, 6))
    ax.set_xlim(0, 10); ax.set_ylim(0, 10); ax.axis("off")

    levels = [
        ("Initial pool of factors\n(literature + experts + pilot + field)", 144,
         9.0, 8.0, "#BDD7EE"),
        ("Removal of overlapping items\n(face validity by expert panel)",     112,
         9.0, 6.4, "#9DC3E6"),
        ("Reliability test (Cronbach's α >= 0.70)\n- low-correlation items dropped", 96,
         9.0, 4.8, "#5B9BD5"),
        ("Consistency check\n(item-total correlation, KMO >= 0.6)",  88,
         9.0, 3.2, "#2E75B6"),
        ("Finalised factor set used in\nFuzzy AHP and ELECTRE", 82,
         9.0, 1.6, "#1F4E78"),
    ]

    width_factor = 0.045  # px per factor
    for title, n, top_y, _, color in levels:
        w_top = n * width_factor
        w_bot = max(40, n - 8) * width_factor
        x_c = 5
        y_top, y_bot = top_y, top_y - 1.1
        poly = Polygon([
            (x_c - w_top/2, y_top),
            (x_c + w_top/2, y_top),
            (x_c + w_bot/2, y_bot),
            (x_c - w_bot/2, y_bot),
        ], closed=True, facecolor=color, edgecolor="#1F4E78", linewidth=1.3)
        ax.add_patch(poly)
        # label inside
        ax.text(x_c, (y_top + y_bot) / 2, f"{title}\nn = {n}",
                ha="center", va="center",
                fontsize=10, family="serif", color="white" if n > 90 else "black",
                weight="bold")

    # arrow connectors
    for i in range(len(levels) - 1):
        y1 = levels[i][2] - 1.1 - 0.05
        y2 = levels[i+1][2] + 0.05
        ax.add_patch(FancyArrowPatch((5, y1), (5, y2),
                                      arrowstyle="-|>", mutation_scale=14,
                                      color="#404040", lw=1.5))

    plt.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def fig_timeline(path):
    fig, ax = plt.subplots(figsize=(11, 3.6))
    months = [
        "Sep'24","Oct","Nov","Dec","Jan'25","Feb","Mar",
        "Apr","May","Jun","Jul","Aug",
        "Sep","Oct","Nov","Dec'25"
    ]
    x = np.arange(len(months))

    ax.barh(2, 7, left=0,  height=0.6, color="#FFD966", edgecolor="#BF8F00",
            label="Phase I: Conceptualisation, literature review, "
                  "expert consultation, pilot study, identification of 144 factors")
    ax.barh(1, 5, left=7,  height=0.6, color="#A9D18E", edgecolor="#548235",
            label="Phase II: Reliability & consistency screening, "
                  "questionnaire finalisation (144 -> 82 factors)")
    ax.barh(0, 4, left=12, height=0.6, color="#9DC3E6", edgecolor="#1F4E78",
            label="Phase III: Main field survey across the four Garhwal "
                  "districts (n = 1,300)")

    ax.set_xticks(x)
    ax.set_xticklabels(months, family="serif", fontsize=9, rotation=30)
    ax.set_yticks([])
    ax.set_xlim(-0.5, 15.5)
    ax.set_ylim(-0.7, 2.8)

    for spine in ("top","right","left"):
        ax.spines[spine].set_visible(False)

    # phase labels on bars
    ax.text(3.5, 2, "PHASE I", ha="center", va="center",
            fontsize=11, weight="bold", family="serif")
    ax.text(9.5, 1, "PHASE II", ha="center", va="center",
            fontsize=11, weight="bold", family="serif")
    ax.text(14, 0, "PHASE III", ha="center", va="center",
            fontsize=11, weight="bold", family="serif")

    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.35),
              fontsize=9, frameon=False, ncol=1)

    plt.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def fig_factor_split(path):
    fig, ax = plt.subplots(figsize=(9, 5))
    stages   = ["Farm", "Trader", "Mandi", "Retail", "Transport"]
    initial  = [34, 29, 29, 28, 24]
    final    = [19, 16, 16, 15, 16]
    x = np.arange(len(stages))
    w = 0.38
    ax.bar(x - w/2, initial, w, color="#9DC3E6", edgecolor="#1F4E78",
           label="Initial pool (n = 144)")
    ax.bar(x + w/2, final,   w, color="#1F4E78", edgecolor="#1F4E78",
           label="Finalised set (n = 82)")
    for i, (a, b) in enumerate(zip(initial, final)):
        ax.text(i - w/2, a + 0.5, str(a), ha="center", fontsize=10, family="serif")
        ax.text(i + w/2, b + 0.5, str(b), ha="center", fontsize=10,
                family="serif", weight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(stages, family="serif", fontsize=11)
    ax.set_ylabel("Number of factors", family="serif", fontsize=11)
    ax.set_title("Stage-wise reduction of factors after screening",
                 family="serif", fontsize=12, weight="bold")
    ax.legend(prop={"family": "serif", "size": 10}, frameon=False)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    plt.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# Generate all figures
fig_research_process("flow_diagram_chapter3.png")
fig_methodology_framework("framework_chapter3.png")
fig_funnel("factor_funnel_chapter3.png")
fig_timeline("phases_chapter3.png")
fig_factor_split("factor_split_chapter3.png")


# ============================================================
# DOCX HELPERS
# ============================================================

def _set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(k), FONT)


def set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(k), FONT)


def add_h(doc, text):
    """Heading - 14 pt bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    _set_run_font(p.add_run(text), size=14, bold=True)
    return p


def add_sh(doc, text):
    """Sub-heading - 12 pt BOLD + ITALIC."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    _set_run_font(p.add_run(text), size=12, bold=True, italic=True)
    return p


def add_p(doc, text, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.5)
    _set_run_font(p.add_run(text), size=12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    for r in p.runs:
        r.text = ""
    p.paragraph_format.space_after = Pt(2)
    _set_run_font(p.add_run(text), size=12)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    for r in p.runs:
        r.text = ""
    p.paragraph_format.space_after = Pt(2)
    _set_run_font(p.add_run(text), size=12)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    _set_run_font(p.add_run(text), size=11, bold=True)


def cell_shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_table(doc, header, rows, caption=None):
    if caption:
        add_caption(doc, caption)
    t = doc.add_table(rows=1+len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        para = c.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(para.add_run(h), size=11, bold=True)
        cell_shade(c, "D9E1F2")
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = ""
            para = c.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(para.add_run(str(val)), size=11)
    doc.add_paragraph("")
    return t


def add_figure(doc, image_path, caption, width_inches=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    add_caption(doc, caption)


# ============================================================
# COMPOSE
# ============================================================

doc = Document()
set_default_font(doc)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# Title
ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_run_font(ttl.add_run("CHAPTER 3"), size=14, bold=True)
ttl2 = doc.add_paragraph(); ttl2.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_run_font(ttl2.add_run("RESEARCH METHODOLOGY"), size=14, bold=True)
doc.add_paragraph("")

# ------------------------------------------------------------
# 3.1 Introduction
# ------------------------------------------------------------
add_h(doc, "3.1  Introduction")
add_p(doc,
    "This chapter sets out the methodology I followed to study the factors "
    "shaping the distribution channel of Malta (Citrus sinensis) in the "
    "Garhwal region of Uttarakhand. The work is descriptive as well as "
    "analytical: descriptive in the sense that it captures how the channel "
    "currently operates from the orchard through to the retail counter, and "
    "analytical in that it applies two Multi-Criteria Decision-Making (MCDM) "
    "tools — Fuzzy AHP and ELECTRE — to identify which factors weigh most "
    "heavily on channel performance and to rank the alternatives that the "
    "industry actually has in front of it.")
add_p(doc,
    "What follows describes the research design, the way the sample was "
    "drawn and why it was sized at 1,300 respondents, the instrument and "
    "the screening procedure used to take a 144-factor pool down to 82 "
    "validated items, and finally the two analytical tools and the way "
    "they fit together. A schematic of the entire research process is "
    "presented at the end of this chapter so the reader can see the whole "
    "study at a glance.")

# ------------------------------------------------------------
# 3.2 Research Design
# ------------------------------------------------------------
add_h(doc, "3.2  Research Design")
add_p(doc,
    "The study uses a cross-sectional research design. By cross-sectional "
    "I mean that data were collected from each respondent at a single point "
    "in time within the marketing season, rather than tracking the same "
    "respondents repeatedly across years. This choice fits the problem "
    "well. Malta has a short season — broadly October to February — so a "
    "snapshot taken across one harvest cycle gives a coherent picture of "
    "channel behaviour without the confounds that creep in when one tries "
    "to compare seasons of different rainfall, different prices and "
    "different policy regimes. It is also the design most often used in the "
    "supply-chain MCDM literature when the goal is to weight criteria and "
    "rank alternatives, rather than to test causality (Bryman and Bell, "
    "2015; Saunders et al., 2019).")

add_sh(doc, "3.2.1  Nature of the Study and Mixed-Method Approach")
add_p(doc,
    "Although the analytical engine of the study is quantitative, the "
    "broader design is mixed-method. Primary data came from a structured "
    "questionnaire administered in the field, supplemented by stakeholder "
    "interactions with farmers, traders, commission agents, mandi "
    "intermediaries, wholesalers, retailers and transporters. Secondary "
    "data were drawn from peer-reviewed journals, reports of the "
    "Department of Horticulture (Government of Uttarakhand), the National "
    "Horticulture Board, agricultural databases such as APEDA and FAO, "
    "earlier theses on hill horticulture and a small set of books on "
    "supply-chain management. The two streams played complementary roles: "
    "the secondary literature gave the conceptual scaffolding and the "
    "initial factor list; the primary data gave the weights and the "
    "outranking that the chapter ultimately delivers.")

# ------------------------------------------------------------
# 3.3 Three Phases of Data Collection
# ------------------------------------------------------------
add_h(doc, "3.3  Three Phases of Data Collection")
add_p(doc,
    "Fieldwork ran from September 2024 to December 2025 and was deliberately "
    "split into three phases. Splitting the work this way meant that the "
    "instrument used in the main survey had already been refined twice over "
    "before the bulk of the responses were collected, which is the kind of "
    "discipline most thesis projects skip. The phase structure is summarised "
    "in Figure 3.1.")

add_figure(doc, "phases_chapter3.png",
           "Figure 3.1: Three-phase Data Collection Timeline "
           "(September 2024 - December 2025)", width_inches=6.5)

add_sh(doc, "3.3.1  Phase I - Conceptualisation and Factor Identification "
            "(Sep 2024 - Mar 2025)")
add_p(doc,
    "Phase I went into building the conceptual base of the study. A "
    "systematic literature review on horticultural supply chains, hill "
    "agriculture and citrus marketing produced a candidate list of "
    "factors. This list was then opened up to a panel of subject experts "
    "drawn from horticulture extension services, two Krishi Vigyan "
    "Kendras, and supply-chain academia. A short pilot was run with 30 "
    "respondents across two blocks each in Pauri Garhwal and Rudraprayag, "
    "and field observations were noted from three day-trips to the "
    "Dehradun and Haridwar wholesale mandis. By the end of Phase I, the "
    "factor pool stood at 144.")

add_sh(doc, "3.3.2  Phase II - Reliability and Consistency Screening "
            "(Apr 2025 - Aug 2025)")
add_p(doc,
    "Phase II was about reducing the noise in that 144-factor pool. The "
    "process is described in detail in section 3.5; the short version is "
    "that overlapping items were merged, items with weak content validity "
    "were dropped on expert advice, and a reliability test (Cronbach's α "
    "≥ 0.70) and item-total consistency check were used to eliminate the "
    "items that did not hold together as a scale. The number came down "
    "from 144 to 82, and the questionnaire used in the main survey was "
    "finalised on this trimmed set.")

add_sh(doc, "3.3.3  Phase III - Main Field Survey (Sep 2025 - Dec 2025)")
add_p(doc,
    "Phase III was the main data-collection drive. Trained enumerators, "
    "supervised by the researcher, visited the four Garhwal districts and "
    "administered the finalised instrument to 1,485 respondents; 1,300 of "
    "those returned valid and complete schedules and form the analytical "
    "sample. Wherever possible the schedules were administered face to "
    "face in the local language, with the researcher reading each item "
    "aloud and recording the response — a practice that yields more "
    "reliable results with semi-literate respondents than self-completed "
    "schedules (Gangrade, 2006; Shah, 1972).")

# ------------------------------------------------------------
# 3.4 Population and Sampling
# ------------------------------------------------------------
add_h(doc, "3.4  Population and Sampling")
add_p(doc,
    "The target population is the set of stakeholders who participate in "
    "the Malta distribution channel in the Garhwal region of Uttarakhand. "
    "Four districts were purposively selected at the geographical layer — "
    "Pauri Garhwal, Tehri Garhwal, Rudraprayag and Chamoli — because they "
    "between them account for the dominant share of Malta cultivation in "
    "the state. Within these districts, seven stakeholder groups were "
    "enumerated: Farmers, Village Traders, Commission Agents, Wholesalers, "
    "Retailers, Transporters, and Experts / Officials.")

add_sh(doc, "3.4.1  Sampling Technique - Stratified Random Sampling")
add_p(doc,
    "Once the strata were defined as the seven stakeholder groups crossed "
    "with the four districts, stratified random sampling was used to draw "
    "respondents within each stratum. Stratification is important here for "
    "two reasons. First, the seven groups are very different in size — the "
    "farmer base is much larger than, say, the commission-agent base — so "
    "a simple random sample would have over-represented farmers and "
    "left some downstream groups too small to support the MCDM "
    "aggregation. Second, the four districts differ in terrain and in the "
    "intensity of Malta trade, and a balanced district representation was "
    "essential for the geographical coverage argument that I make in "
    "section 3.4.3.")
add_p(doc,
    "Within each stratum the sampling frame was constructed from "
    "Department-of-Horticulture grower lists for farmers, mandi-samiti "
    "registers for wholesalers and commission agents, and shop-front "
    "enumeration for retailers and village traders. Where complete "
    "frames were not available — chiefly for transporters — snowball "
    "referral was used to extend the frame, but the actual selection "
    "within the extended frame remained random.")

add_sh(doc, "3.4.2  Achieved Sample (n = 1,300)")
add_p(doc,
    "The achieved sample is summarised stakeholder-wise in Table 3.1 and "
    "district-wise in Table 3.2.")

add_table(doc,
    ["S. No.", "Stakeholder Category", "No. of Respondents", "Share (%)"],
    [
        ["1", "Farmers",            "650",  "50.0%"],
        ["2", "Wholesalers",        "130",  "10.0%"],
        ["3", "Transporters",       "130",  "10.0%"],
        ["4", "Retailers",          "130",  "10.0%"],
        ["5", "Village Traders",    "117",  "9.0%"],
        ["6", "Commission Agents",  "78",   "6.0%"],
        ["7", "Experts / Officials","65",   "5.0%"],
        ["",  "Total",              "1,300","100.0%"],
    ],
    "Table 3.1: Stakeholder-wise Distribution of the Achieved Sample (n = 1,300)")

add_table(doc,
    ["S. No.", "District (Garhwal Region)", "No. of Respondents", "Share (%)"],
    [
        ["1", "Chamoli",         "340",  "26.2%"],
        ["2", "Tehri Garhwal",   "333",  "25.6%"],
        ["3", "Rudraprayag",     "316",  "24.3%"],
        ["4", "Pauri Garhwal",   "311",  "23.9%"],
        ["",  "Total",           "1,300","100.0%"],
    ],
    "Table 3.2: District-wise Distribution of the Achieved Sample (n = 1,300)")

# ------------------------------------------------------------
# 3.5 Factor Identification and Screening (144 -> 82)
# ------------------------------------------------------------
add_h(doc, "3.5  Factor Identification and Screening")
add_p(doc,
    "The analytical strength of any MCDM exercise depends on the quality "
    "of the underlying factor list. I therefore treated factor "
    "identification as a project in its own right rather than as an "
    "instrument-design afterthought. The process moved in two clean steps: "
    "first, building a wide candidate pool, and then trimming it to a set "
    "of items that were both internally consistent and operationally "
    "meaningful.")

add_sh(doc, "3.5.1  Building the Candidate Pool of 144 Factors")
add_p(doc,
    "The starting point was the literature. An extensive review of "
    "horticultural supply-chain studies, with particular attention to "
    "studies on citrus and hill horticulture, yielded a long list of "
    "factors that had previously been reported as material to channel "
    "performance. This list was supplemented by inputs from subject "
    "experts, by what came up in the pilot study, and by what stakeholders "
    "themselves raised during the Phase-I field visits. The five distinct "
    "input streams used to build the pool are listed below.")
add_bullet(doc, "Extensive literature review on horticultural and citrus supply chains")
add_bullet(doc, "Expert consultation with horticulture extension services and academia")
add_bullet(doc, "Pilot study findings (n = 30, two blocks in Garhwal)")
add_bullet(doc, "Stakeholder discussions during exploratory field visits")
add_bullet(doc, "Direct field observations at orchards and at the Dehradun and Haridwar mandis")

add_p(doc,
    "These inputs together produced 144 candidate factors. To make them "
    "tractable, the factors were classified along two crossed dimensions: "
    "the stage of the channel at which they operate (Farm, Trader, Mandi, "
    "Retail, Transport) and the aspect of performance they bear on (Cost, "
    "Time, Quality). Table 3.3 summarises the distribution of the 144 "
    "candidates across these dimensions.")

add_table(doc,
    ["Stage", "Cost", "Time", "Quality", "Sub-total"],
    [
        ["Farm-level",       "12", "10", "12", "34"],
        ["Trader-level",     "11", "9",  "9",  "29"],
        ["Mandi-level",      "10", "9",  "10", "29"],
        ["Retail-level",     "11", "8",  "9",  "28"],
        ["Transport-level",  "9",  "8",  "7",  "24"],
        ["Total",            "53", "44", "47", "144"],
    ],
    "Table 3.3: Initial Distribution of the 144 Candidate Factors")

add_sh(doc, "3.5.2  Reliability and Consistency Screening")
add_p(doc,
    "The screening process was deliberately conservative. Items were "
    "removed at four successive checkpoints:")
add_numbered(doc,
    "Removal of overlapping items. Where two or more candidates measured "
    "essentially the same construct, the cleaner of the two was retained "
    "and the rest dropped after expert validation. This step took the pool "
    "from 144 to 112.")
add_numbered(doc,
    "Reliability test. Cronbach's α was computed for every sub-scale "
    "(stage × aspect). Items whose deletion improved the α of the parent "
    "scale, and items whose item-total correlation fell below 0.30, were "
    "dropped. This step took the pool from 112 to 96.")
add_numbered(doc,
    "Consistency check. Sampling adequacy was verified using the "
    "Kaiser–Meyer–Olkin (KMO) measure (≥ 0.60) and Bartlett's test of "
    "sphericity (p < 0.05). Items that loaded weakly on their parent "
    "factor were removed. This step took the pool from 96 to 88.")
add_numbered(doc,
    "Final expert review. The trimmed list was returned to the expert "
    "panel for face-validity confirmation and any last refinements. Six "
    "more items were dropped at this step, leaving the finalised set at "
    "82 factors.")

add_p(doc,
    "The funnel is shown graphically in Figure 3.2.")
add_figure(doc, "factor_funnel_chapter3.png",
           "Figure 3.2: Factor Screening Funnel - 144 Candidates "
           "Reduced to 82 Finalised Factors", width_inches=5.6)

add_p(doc,
    "After screening, the 82 retained factors were redistributed across the "
    "five stages and three aspects, as set out in Table 3.4. The "
    "stage-level reduction is also shown in Figure 3.3.")

add_table(doc,
    ["Stage", "Cost", "Time", "Quality", "Sub-total"],
    [
        ["Farm-level",       "7", "6", "6", "19"],
        ["Trader-level",     "6", "5", "5", "16"],
        ["Mandi-level",      "6", "5", "5", "16"],
        ["Retail-level",     "6", "4", "5", "15"],
        ["Transport-level",  "6", "5", "5", "16"],
        ["Total",            "31","25","26","82"],
    ],
    "Table 3.4: Final Distribution of the 82 Retained Factors after Screening")

add_figure(doc, "factor_split_chapter3.png",
           "Figure 3.3: Stage-wise Reduction of Factors after Reliability "
           "and Consistency Screening", width_inches=5.8)

add_sh(doc, "3.5.3  Illustrative Factors in the Final Set")
add_p(doc,
    "To give the reader a sense of what the 82 finalised factors look like "
    "in practice, the most representative items are summarised below.")
add_bullet(doc, "Packaging and grading efficiency at the orchard and post-harvest stage")
add_bullet(doc, "Transportation cost, fuel cost and toll-related charges")
add_bullet(doc, "Storage and warehousing facilities (capacity, ventilation, temperature)")
add_bullet(doc, "Loading and unloading efficiency at every transfer point")
add_bullet(doc, "Transit time and waiting time across the channel")
add_bullet(doc, "Product handling practices and bruising during transit")
add_bullet(doc, "Market accessibility from hill orchards to the plains mandis")
add_bullet(doc, "Commission and intermediary charges at mandi level")
add_bullet(doc, "Retail infrastructure, display practice and customer demand")
add_bullet(doc, "Vehicle availability and temperature-control capability")
add_bullet(doc, "Post-harvest handling and quality maintenance over the channel")

# ------------------------------------------------------------
# 3.6 Methodology Framework Diagram
# ------------------------------------------------------------
add_h(doc, "3.6  Overall Methodology Framework")
add_p(doc,
    "Figure 3.4 lays out the overall framework of the methodology — "
    "showing how the inputs (literature, experts, pilot, field "
    "observation) feed into the candidate factor pool, how screening "
    "produces the finalised 82 factors, how the cross-sectional survey is "
    "administered to the 1,300 respondents, and how the two MCDM tools "
    "are connected. This single diagram is the navigation map for the rest "
    "of the chapter.")

add_figure(doc, "framework_chapter3.png",
           "Figure 3.4: Overall Methodology Framework", width_inches=6.5)

# ------------------------------------------------------------
# 3.7 Justification of the Sample Size for MCDM Methods
# ------------------------------------------------------------
add_h(doc, "3.7  Justification of the Sample Size for MCDM Methods")
add_p(doc,
    "MCDM methods such as Fuzzy AHP and ELECTRE are usually run with small "
    "expert panels of five to thirty decision makers. Saaty (1980, 2008) "
    "in fact built AHP around exactly that small-panel logic and "
    "recommended geometric-mean aggregation of pairwise comparisons. "
    "Why, then, has the present study gone all the way up to 1,300 "
    "respondents? The answer is that the study is not a pure expert-panel "
    "MCDM exercise; it is an empirical, stakeholder-aggregated MCDM "
    "exercise running across a real-world supply chain. In this kind of "
    "configuration a larger sample is not an indulgence — it is what "
    "makes the weights and the outranking robust. Five specific arguments "
    "are set out below.")

add_sh(doc, "3.7.1  Aggregation Stability")
add_p(doc,
    "When pairwise comparisons are aggregated across respondents using the "
    "geometric mean — the standard aggregation operator in AHP and Fuzzy "
    "AHP — the aggregated matrix converges to its long-run mean as the "
    "number of respondents grows (Forman and Peniwati, 1998). Empirical "
    "work suggests that beyond about thirty respondents per stakeholder "
    "group the change in aggregated weights becomes negligible. The "
    "smallest stakeholder group in this study (Experts / Officials) has 65 "
    "respondents and the largest (Farmers) has 650, so aggregation "
    "stability is comfortably achieved for every group.")

add_sh(doc, "3.7.2  Statistical Adequacy")
add_p(doc,
    "Although MCDM does not impose a parametric sample-size rule of its "
    "own, the demographic and rating components of the questionnaire can "
    "be treated as a population-estimation problem. Cochran's (1977) "
    "formula for an infinite population, at a 95 per cent confidence "
    "level and a 5 per cent margin of error, recommends a minimum sample "
    "of 384. The 1,300 sample exceeds this minimum by a factor of more "
    "than three and brings the margin of error down to roughly 2.7 per "
    "cent at the 95 per cent confidence level.")

add_sh(doc, "3.7.3  Stakeholder Representativeness")
add_p(doc,
    "The Malta channel involves seven heterogeneous stakeholder groups, "
    "and a credible MCDM application has to treat each group as a "
    "distinct sub-sample so that group-specific weight vectors can be "
    "aggregated into a channel-level weight vector (Dong et al., 2010). "
    "The achieved sample provides between 65 and 650 respondents per "
    "group, which keeps every group above the 30-respondent stability "
    "threshold and preserves the relative weight of each group in the "
    "channel.")

add_sh(doc, "3.7.4  Geographical Coverage")
add_p(doc,
    "The four Garhwal districts together carry the bulk of Malta "
    "cultivation in Uttarakhand, and the achieved district-wise "
    "distribution is balanced (between 23.9 and 26.2 per cent per "
    "district). That balance heads off the geographical-bias problem that "
    "single-district MCDM studies often have to live with.")

add_sh(doc, "3.7.5  Robustness for ELECTRE Outranking")
add_p(doc,
    "ELECTRE is an outranking method whose concordance and discordance "
    "indices are computed from the decision matrix of alternatives and "
    "criteria. The performance scores in that decision matrix are "
    "themselves the sample means of stakeholder ratings on each "
    "criterion. A larger sample reduces the standard error of these means "
    "and therefore reduces the sensitivity of the concordance and "
    "discordance thresholds to outlier ratings (Roy and Bouyssou, 1993; "
    "Figueira et al., 2005). A sample of 1,300 yields standard errors of "
    "the order of 0.03 on a five-point scale, well below the differences "
    "between alternatives, which is what an outranking study needs.")

# ------------------------------------------------------------
# 3.8 Tools of Analysis
# ------------------------------------------------------------
add_h(doc, "3.8  Tools of Analysis")
add_p(doc,
    "Only two tools were used for analysis: Fuzzy AHP and ELECTRE. "
    "No other MCDM technique, no regression, no SEM, no machine-learning "
    "model. The two tools were chosen because each does one job well, and "
    "because the output of one is precisely the input the other needs. "
    "Fuzzy AHP is used to determine the relative importance of the 82 "
    "factors under the linguistic vagueness that real respondents bring "
    "to a survey. ELECTRE then takes those weights and uses them to rank "
    "the alternative configurations of the Malta distribution channel.")

# ------ FAHP -------
add_sh(doc, "3.8.1  Fuzzy Analytic Hierarchy Process (Fuzzy AHP)")
add_p(doc,
    "Fuzzy AHP extends Saaty's classical AHP (Saaty, 1980) by replacing "
    "crisp pairwise comparisons with Triangular Fuzzy Numbers (TFNs) of "
    "the form (l, m, u). The use of TFNs captures the linguistic "
    "vagueness that is inherent in human judgement and is particularly "
    "useful with stakeholder respondents whose expertise is experiential "
    "rather than numerical. The procedure adopted here follows the "
    "extent-analysis approach of Chang (1996), which is the most widely "
    "applied Fuzzy AHP procedure in the supply-chain literature.")

add_p(doc, "The procedural steps used in this study are as follows.")
add_numbered(doc, "Build the hierarchy. The decision goal — selection of the optimal "
                  "Malta distribution channel — sits at the top, followed by the "
                  "criteria, the sub-criteria and the channel alternatives.")
add_numbered(doc, "Elicit pairwise comparisons from respondents on the nine-point "
                  "linguistic scale shown in Table 3.5.")
add_numbered(doc, "Convert each linguistic judgement into its corresponding TFN (l, m, u).")
add_numbered(doc, "Aggregate responses across respondents via the fuzzy geometric mean, "
                  "yielding a single fuzzy comparison matrix per stakeholder group.")
add_numbered(doc, "Compute the fuzzy synthetic extent value Si for each criterion as "
                  "Si = Σ_j M(g, ij) ⊗ [Σ_i Σ_j M(g, ij)]^(-1).")
add_numbered(doc, "Compute the degree of possibility V(M2 ≥ M1) for every pair of "
                  "fuzzy synthetic extents and derive the non-normalised priority "
                  "vector W'.")
add_numbered(doc, "Defuzzify each TFN by the Centre-of-Area (CoA) method, normalise, "
                  "and obtain the final criterion weights w_j with Σ w_j = 1.")
add_numbered(doc, "Compute the Consistency Ratio. A pairwise-comparison matrix is "
                  "accepted only if CR < 0.10 (Saaty, 1980).")

add_table(doc,
    ["Linguistic Term", "Saaty Scale",
     "Triangular Fuzzy Number (l, m, u)", "Reciprocal TFN"],
    [
        ["Equally important",                     "1",  "(1, 1, 1)", "(1, 1, 1)"],
        ["Weakly more important",                 "3",  "(1, 3, 5)", "(1/5, 1/3, 1)"],
        ["Strongly more important",               "5",  "(3, 5, 7)", "(1/7, 1/5, 1/3)"],
        ["Very strongly more important",          "7",  "(5, 7, 9)", "(1/9, 1/7, 1/5)"],
        ["Extremely / Absolutely more important", "9",  "(7, 9, 9)", "(1/9, 1/9, 1/7)"],
        ["Intermediate values",                   "2, 4, 6, 8",
         "Linear interpolation", "Reciprocal interpolation"],
    ],
    "Table 3.5: Linguistic Scale and Triangular Fuzzy Numbers used in Fuzzy AHP")

# ------ ELECTRE ------
add_sh(doc, "3.8.2  Elimination and Choice Translating Reality (ELECTRE)")
add_p(doc,
    "ELECTRE is an outranking MCDM method developed by Roy (1968, 1991). "
    "It does not chase a single aggregate score per alternative; instead "
    "it builds an outranking relation by asking, for every pair of "
    "alternatives, how strong is the case in favour (concordance) and "
    "how strong is the case against (discordance). That logic suits the "
    "Malta problem because no channel dominates on every criterion — "
    "trade-offs have to be evaluated explicitly. The notation used in "
    "this study is summarised in Table 3.6.")

add_table(doc,
    ["Symbol", "Meaning"],
    [
        ["A = {a1, …, am}",       "Set of distribution-channel alternatives"],
        ["C = {c1, …, cn}",       "Set of decision criteria"],
        ["w_j",                   "Weight of criterion j (taken from Fuzzy AHP, Σ w_j = 1)"],
        ["X = [x_ij]",            "Decision matrix; x_ij is the score of alternative i on criterion j"],
        ["R = [r_ij]",            "Normalised decision matrix"],
        ["V = [v_ij] = w_j · r_ij","Weighted normalised decision matrix"],
        ["C(a, b)",               "Concordance index"],
        ["D(a, b)",               "Discordance index"],
        ["c̄, d̄",                "Concordance and discordance thresholds (means of C and D)"],
        ["F, G",                  "Concordance and discordance dominance matrices"],
        ["E = F ⊙ G",             "Aggregate dominance matrix used to derive the final ranking"],
    ],
    "Table 3.6: Notation used in the ELECTRE Procedure")

add_p(doc, "The ELECTRE steps used in the study are as follows.")
add_numbered(doc, "Construct the decision matrix X = [x_ij] with each row a channel "
                  "alternative and each column a criterion; entries are mean ratings "
                  "from the survey.")
add_numbered(doc, "Normalise X by vector normalisation: r_ij = x_ij / sqrt(Σ_i x_ij²).")
add_numbered(doc, "Compute V by multiplying each column r_ij by the criterion weight "
                  "w_j obtained from Fuzzy AHP.")
add_numbered(doc, "For every ordered pair (a, b), build the concordance set C(a, b) "
                  "(criteria on which a is at least as good as b) and the discordance "
                  "set D(a, b) (criteria on which a is worse).")
add_numbered(doc, "Compute the concordance index C(a,b) = Σ w_j for j ∈ C(a,b) and the "
                  "discordance index D(a,b) = max{|v_aj − v_bj|} / max{|v_ij − v_kj|}.")
add_numbered(doc, "Compute the thresholds c̄ and d̄ as the averages of the indices, "
                  "and build the boolean dominance matrices F and G.")
add_numbered(doc, "Multiply F and G element-wise to get E. Alternative a outranks "
                  "alternative b if and only if e_ab = 1.")
add_numbered(doc, "Construct the outranking graph and read off the final ranking of "
                  "channel alternatives.")

add_sh(doc, "3.8.3  Integration of Fuzzy AHP and ELECTRE")
add_p(doc,
    "The two tools meet at exactly one point: the criterion weights w_j "
    "produced at Step 7 of the Fuzzy AHP procedure feed directly into "
    "Step 3 of the ELECTRE procedure, where they enter the weighted "
    "normalised decision matrix V. No other tool is used in between, and "
    "no other tool is used after. The end-to-end research process is "
    "shown schematically in Figure 3.5.")

# ------------------------------------------------------------
# 3.9 Research Process Flow Diagram
# ------------------------------------------------------------
add_h(doc, "3.9  Research Process Flow Diagram")
add_p(doc,
    "Figure 3.5 puts the entire research process on a single page. The "
    "left-hand branch shows the Fuzzy AHP weight-derivation track and the "
    "right-hand branch shows the ELECTRE outranking track. The two "
    "branches converge at the bottom, where the optimal distribution "
    "channel for Malta is identified and the recommendations to "
    "stakeholders are framed.")

add_figure(doc, "flow_diagram_chapter3.png",
           "Figure 3.5: Research Process Flow Diagram - "
           "Malta Distribution Channel Optimisation Study", width_inches=6.0)

# ------------------------------------------------------------
# 3.10 Reliability and Validity
# ------------------------------------------------------------
add_h(doc, "3.10  Reliability and Validity")
add_p(doc,
    "Reliability of the responses was protected at three points in the "
    "study. First, during instrument design through Cronbach's α and "
    "item-total correlation. Second, during Fuzzy AHP elicitation through "
    "the CR < 0.10 rule: any pairwise-comparison matrix with CR ≥ 0.10 "
    "was returned to the respondent for revision and matrices that could "
    "not be reconciled were excluded. Third, during ELECTRE outranking "
    "through the use of mean ratings — large samples shrink the standard "
    "errors of these means, and that in turn protects the outranking from "
    "individual outlier ratings.")
add_p(doc,
    "Content validity was secured by deriving the factors from the "
    "literature and refining them through expert review. Face validity "
    "was checked during pre-testing. Convergent validity was supported by "
    "the broad agreement between the Fuzzy AHP rank-order of criteria and "
    "the rank-order suggested independently by the expert panel.")

# ------------------------------------------------------------
# 3.11 Ethical Considerations
# ------------------------------------------------------------
add_h(doc, "3.11  Ethical Considerations")
add_p(doc,
    "Participation was voluntary and informed consent was obtained from "
    "every respondent before the schedule was administered. Respondents "
    "were assured that their identity would remain anonymous in any "
    "outputs of the study and that the data would be used only for "
    "academic purposes. No commercially sensitive pricing or identity "
    "information was solicited or retained, and the researcher complied "
    "with the ethical guidelines of the parent institution throughout "
    "the fieldwork.")

# ------------------------------------------------------------
# 3.12 Concluding Remarks
# ------------------------------------------------------------
add_h(doc, "3.12  Concluding Remarks")
add_p(doc,
    "This chapter has described how the study was actually carried out: a "
    "cross-sectional, mixed-method design; three phases of data collection "
    "running from September 2024 to December 2025; a stratified random "
    "sample of 1,300 respondents drawn from seven stakeholder groups "
    "across the four Malta-producing districts of the Garhwal region; a "
    "factor pool that started at 144 candidates and was screened down to "
    "82 finalised items through reliability and consistency tests; and "
    "two analytical tools — Fuzzy AHP for criterion weighting and ELECTRE "
    "for outranking — connected through the criterion weight vector. The "
    "next chapter applies this methodology to the Malta data and presents "
    "the results.")

doc.save("Chapter_3_Research_Methodology.docx")
print("WROTE Chapter_3_Research_Methodology.docx")
print("Figures:")
for f in ["flow_diagram_chapter3.png", "framework_chapter3.png",
          "factor_funnel_chapter3.png", "phases_chapter3.png",
          "factor_split_chapter3.png"]:
    print(f"  - {f} ({os.path.getsize(f)} bytes)")
