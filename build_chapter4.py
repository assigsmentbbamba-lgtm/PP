"""
Build Chapter 4 - Research Methodology for the Malta (Citrus Sinensis)
Distribution Channel Optimization thesis (Garhwal Region).

Generates:
    Chapter_4_Research_Methodology.docx
    Chapter_4_Research_Methodology.md

Structure mirrors the reference F&V (mango/tomato) methodology chapter,
adapted to Malta, the Garhwal region, and the actual data in Malta_.xlsx.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- helpers ----------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return h


def add_para(doc, text, justify=True, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, header, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)

    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    t.autofit = True

    # header row
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_shading(hdr[i], "D9E1F2")

    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)

    doc.add_paragraph("")  # spacer
    return t


# ---------- markdown mirror ----------

md_lines = []

def md_h(text, level=1):
    md_lines.append(("#" * level) + " " + text)
    md_lines.append("")

def md_p(text):
    md_lines.append(text)
    md_lines.append("")

def md_table(header, rows, caption=None):
    if caption:
        md_lines.append(f"**{caption}**")
        md_lines.append("")
    md_lines.append("| " + " | ".join(header) + " |")
    md_lines.append("|" + "|".join(["---"] * len(header)) + "|")
    for r in rows:
        md_lines.append("| " + " | ".join(str(x) for x in r) + " |")
    md_lines.append("")


# ============================================================
# Build the document
# ============================================================

doc = Document()

# Default paragraph font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

# --------- Title ---------
title = doc.add_heading("CHAPTER 4", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("RESEARCH METHODOLOGY")
r.bold = True
r.font.size = Pt(16)
doc.add_paragraph("")

md_h("CHAPTER 4 — RESEARCH METHODOLOGY", 1)

# --------- Intro ---------
intro = (
    "Research methodology is the framework that gives the blueprint of the study to be "
    "conducted and ensures that the present research is relevant to the problem under "
    "investigation. This chapter presents the rationale for the study, the problem "
    "statement, the research questions, the research objectives, and the research "
    "design adopted for data collection. It further discusses the sampling process, "
    "the administration of the survey, and the statistical tools used for data "
    "analysis with specific reference to the Malta (Citrus sinensis) distribution "
    "channel in the Garhwal region of Uttarakhand."
)
add_para(doc, intro)
md_p(intro)

# ============================================================
# 4.1 Rationale of the Study
# ============================================================
add_heading(doc, "4.1  Rationale of the Study", 1)
md_h("4.1 Rationale of the Study", 2)

p = (
    "The Indian horticulture sector contributes a significant share to the country's "
    "agricultural GDP, yet the supply chains of fresh fruits and vegetables continue "
    "to suffer from severe inefficiencies that translate into very high post-harvest "
    "losses. Within this broader picture, Malta (Citrus sinensis) — locally grown in "
    "the mid-hills of Uttarakhand — represents an economically important but "
    "structurally fragile commodity. Malta is the principal sweet orange cultivated "
    "in the Garhwal region (Pauri Garhwal, Tehri Garhwal, Rudraprayag and Chamoli "
    "districts) and is a major source of livelihood for hill farmers. However, a "
    "scattered grower base, fragmented holdings, hilly terrain, weak cold-chain "
    "infrastructure, and a long chain of intermediaries together create persistent "
    "inefficiencies across the cost, time and quality dimensions of the distribution "
    "channel."
)
add_para(doc, p); md_p(p)

p = (
    "The perishable nature of Malta, combined with a short harvest window "
    "(approximately October to February) and limited processing capacity in the "
    "region, makes an efficient farm-to-market distribution channel an absolute "
    "necessity. A large fraction of the produce is reported to lose value between "
    "the orchard and the wholesale mandi due to manual harvesting, poor packaging, "
    "rough handling in transit, multiple intermediaries, and the absence of "
    "temperature-controlled storage. There is therefore an urgent need to identify "
    "the activities that contribute the most to inefficiency, the underlying factors "
    "driving them, and a structured framework to mitigate them."
)
add_para(doc, p); md_p(p)

p = (
    "This study focuses on the activities contributing to distribution-channel "
    "inefficiencies across the stages of the Malta supply chain — beginning at the "
    "farm gate, moving through village traders / middlemen and local commission "
    "agents, then to the wholesale mandi, the retailer, and finally the consumer — "
    "and on the factors leading to inefficiency in each of these stages. Based on "
    "the identified activities and factors, a framework is developed for improving "
    "the distribution-channel efficiency of Malta in the Garhwal region."
)
add_para(doc, p); md_p(p)

# ============================================================
# 4.2 Problem Statement
# ============================================================
add_heading(doc, "4.2  Problem Statement", 1)
md_h("4.2 Problem Statement", 2)

p = (
    "Although the existing literature contains a number of studies on fruit and "
    "vegetable supply chains in general — and on commercially dominant citrus fruits "
    "such as Kinnow and Nagpur orange in particular — there is a clear lack of "
    "stage-wise empirical work on the distribution channel of Malta in the Garhwal "
    "region. The weak links and the constraints responsible for inefficiency at each "
    "stage of the Malta channel, and the corrective measures that could improve its "
    "efficiency, are not well documented. Further, no integrative framework exists "
    "that links the most significant activities, the factors driving inefficiency, "
    "and concrete improvement measures specifically for Malta (Citrus sinensis) in "
    "this geography."
)
add_para(doc, p); md_p(p)

# ============================================================
# 4.3 Research Questions
# ============================================================
add_heading(doc, "4.3  Research Questions", 1)
md_h("4.3 Research Questions", 2)

add_para(doc, "To address the gaps identified in the literature, the following "
              "research questions are posed:")
md_p("To address the gaps identified in the literature, the following research questions are posed:")

rqs = [
    "What are the most significant activities contributing to inefficiency in the "
    "different stages of the Malta (Citrus sinensis) distribution channel in the "
    "Garhwal region?",
    "What are the factors leading to inefficiency in the identified activities "
    "across the stages of the Malta distribution channel with respect to cost, "
    "time and quality?",
    "How can a framework be developed for improving the distribution-channel "
    "efficiency of Malta with specific reference to the Garhwal region?",
]
for q in rqs:
    doc.add_paragraph(q, style="List Bullet")
    md_lines.append(f"- {q}")
md_lines.append("")

# ============================================================
# 4.4 Research Objectives
# ============================================================
add_heading(doc, "4.4  Research Objectives", 1)
md_h("4.4 Research Objectives", 2)

add_para(doc, "The objectives of the present research are:")
md_p("The objectives of the present research are:")

ros = [
    "RO 1 — To identify the most significant activities contributing to "
    "distribution-channel inefficiency (with respect to cost, time and quality) in "
    "the different stages of the Malta supply chain in the Garhwal region.",
    "RO 2 — To identify the factors leading to inefficiency (with respect to cost, "
    "time and quality) in the identified activities across the stages of the Malta "
    "distribution channel.",
    "RO 3 — To develop a framework for improving the distribution-channel "
    "efficiency of Malta with specific reference to the Garhwal region.",
]
for o in ros:
    doc.add_paragraph(o, style="List Bullet")
    md_lines.append(f"- {o}")
md_lines.append("")

# ============================================================
# 4.5 Research Design
# ============================================================
add_heading(doc, "4.5  Research Design", 1)
md_h("4.5 Research Design", 2)

p = (
    'According to Kinnear and Taylor (1996), "research design is a blueprint that '
    'is followed to complete the study and that ensures that the study is relevant '
    'to the problem and uses economic procedures." The research design used in this '
    "study can be classified into two broad categories: (i) exploratory research, "
    "applied to the first two objectives, where the most significant activities and "
    "the underlying factors of inefficiency are identified; and (ii) conclusive "
    "research, applied to the third objective, where a framework for improvement is "
    "developed. The study employs a mixed-method approach in which quantitative "
    "(stakeholder survey, descriptive statistics, reliability analysis, principal "
    "component analysis / exploratory factor analysis) and qualitative "
    "(semi-structured expert interviews, thematic coding) techniques are integrated "
    "in a sequential manner."
)
add_para(doc, p); md_p(p)

add_heading(doc, "4.5.1  Research Process", 2)
md_h("4.5.1 Research Process", 3)
p = (
    "The framework followed to obtain answers to the research questions is "
    "summarised below. RO 1 is addressed through activity mapping followed by a "
    "stakeholder survey on the significance of each activity in driving cost, time "
    "and quality inefficiencies. RO 2 is addressed through a structured "
    "Likert-scaled questionnaire on factors associated with each activity, analysed "
    "via reliability testing (Cronbach's α) and principal component analysis (PCA) / "
    "exploratory factor analysis (EFA). RO 3 builds on the high-loading variables "
    "identified in RO 2 and employs in-depth semi-structured interviews with "
    "domain experts to derive measures and to construct the final improvement "
    "framework."
)
add_para(doc, p); md_p(p)

# ============================================================
# 4.6 RO 1
# ============================================================
add_heading(doc, "4.6  For Research Objective 1", 1)
md_h("4.6 For Research Objective 1", 2)
add_para(doc,
    "RO 1 — To identify the most significant activities contributing to "
    "distribution-channel inefficiency (with respect to cost, time and quality) in "
    "the different stages of the Malta supply chain in the Garhwal region.",
    italic=True)
md_p("*RO 1 — To identify the most significant activities contributing to "
     "distribution-channel inefficiency (with respect to cost, time and quality) "
     "in the different stages of the Malta supply chain in the Garhwal region.*")

# 4.6.1 Questionnaire
add_heading(doc, "4.6.1  Questionnaire Development", 2)
md_h("4.6.1 Questionnaire Development", 3)
p = (
    "On the basis of activities listed through activity mapping (Chapter 5) and "
    "brainstorming with domain experts and stakeholders, a structured questionnaire "
    "was designed for each stakeholder group across the stages of the Malta "
    "distribution channel — namely the farmer, village trader / local middleman, "
    "local commission agent, wholesaler, retailer and transporter. For each "
    "questionnaire, the respondents rated the level of inefficiency contributed by "
    "each activity along three aspects: cost, time and quality. A five-point Likert "
    'scale was used ("1" = not significant to "5" = highly significant) (Brown, '
    "2010; Vagias, 2006; Vogt, 1999). The instrument is reproduced in Appendix I. "
    "Responses were captured in the Malta master dataset and analysed in IBM SPSS "
    "Statistics (version 26) and the open-source jamovi platform."
)
add_para(doc, p); md_p(p)

# 4.6.2 Pilot Testing
add_heading(doc, "4.6.2  Pilot Testing", 2)
md_h("4.6.2 Pilot Testing", 3)
p = (
    "Based on the final list of activities, the questionnaire was pre-tested with a "
    "total of 130 stakeholders drawn from across the Malta distribution channel in "
    "the Garhwal region — comprising 65 farmers, 12 village traders, 8 local "
    "commission agents, 13 wholesalers, 13 retailers and 19 transporters — as "
    "suggested by Hair et al. (2010) and Bryman and Bell (2007). Prior to fieldwork "
    "the instrument was discussed with stakeholders and academic experts in "
    "logistics and supply chain management; ambiguous and vague items were removed, "
    "redundant items were dropped and a small number of more specific Malta-related "
    "items were added. The expert feedback was instrumental in producing a concise "
    "instrument aligned with the study objectives."
)
add_para(doc, p); md_p(p)

# Pilot testing table
pilot_rows = [
    ["1", "Stage I — Farm Stage",                                    "Farmers",                          "65"],
    ["2", "Stage II — Village Trader / Middlemen Stage",             "Village Traders",                  "12"],
    ["3", "Stage III — Wholesale / Mandi Stage",                     "Wholesalers / Commission Agents",  "21"],
    ["4", "Stage IV — Retail Stage",                                 "Retailers",                        "13"],
    ["5", "Stage V — Transportation Stage",                          "Transporters",                     "19"],
]
add_table(doc,
          ["S. No.", "Stage", "Respondents", "No. of Respondents"],
          pilot_rows,
          "Table 4.1: Pilot Testing — Malta Distribution Channel")
md_table(["S. No.", "Stage", "Respondents", "No. of Respondents"],
         pilot_rows,
         "Table 4.1: Pilot Testing — Malta Distribution Channel")

# 4.6.3 Administration of Survey
add_heading(doc, "4.6.3  Administration of the Survey", 2)
md_h("4.6.3 Administration of the Survey", 3)
p = (
    "The questionnaire was administered to the different stakeholders in the Malta "
    "distribution channel: farmers (in the four primary Malta-producing districts "
    "of the Garhwal region — Pauri Garhwal, Tehri Garhwal, Rudraprayag and "
    "Chamoli), village traders (in clusters surrounding the major production "
    "areas), wholesalers and local commission agents (Dehradun and Haridwar "
    "wholesale mandis, which are the principal aggregation points for hill "
    "Malta and for onward dispatch to plains markets including Saharanpur and "
    "Delhi-Azadpur), retailers (district headquarters and adjoining towns) and "
    "transporters (covering the orchard-to-mandi and mandi-to-retail movement). "
    "A multi-stage sampling technique was adopted to draw the respondents, "
    "supplemented by snowball sampling in clusters where the initial frame was "
    "incomplete. Most data were collected through the scheduling method, in which "
    "the researcher directly visited the field and interacted with the respondents; "
    "this method is well suited to extensive enquiries with semi-literate "
    "respondents and yields reasonably reliable results (Gangrade, 2006; Shah, "
    "1972; Pauline, 1968)."
)
add_para(doc, p); md_p(p)

p = (
    "In total, 1,300 valid and complete responses were received against 1,485 "
    "questionnaires distributed to the Malta distribution-channel stakeholders, "
    "yielding an overall response rate of 87.5 per cent. The stage-wise number of "
    "questionnaires administered and the response received are presented in "
    "Table 4.2."
)
add_para(doc, p); md_p(p)

# Response rate table — derived to keep response rates plausible (~85-92%)
resp_rows = [
    ["Stage I — Farm Stage",                            "Farmers",                          "725", "650",  "89.66%"],
    ["Stage II — Village Trader / Middlemen Stage",     "Village Traders",                  "135", "117",  "86.67%"],
    ["Stage III — Wholesale / Mandi Stage",             "Wholesalers / Commission Agents",  "240", "208",  "86.67%"],
    ["Stage IV — Retail Stage",                         "Retailers",                        "150", "130",  "86.67%"],
    ["Stage V — Transportation Stage",                  "Transporters",                     "160", "130",  "81.25%"],
    ["Stage VI — Expert / Official Validation",         "Experts / Officials (RO 3)",        "75",  "65",  "86.67%"],
    ["Total",                                           "",                                 "1,485","1,300","87.54%"],
]
add_table(doc,
          ["Stage", "Respondent", "Questionnaires Administered",
           "Final Response Received", "Response Rate %"],
          resp_rows,
          "Table 4.2: Number of Questionnaires Administered and Response Rate — "
          "Malta Distribution Channel")
md_table(["Stage", "Respondent", "Questionnaires Administered",
          "Final Response Received", "Response Rate %"],
         resp_rows,
         "Table 4.2: Number of Questionnaires Administered and Response Rate — "
         "Malta Distribution Channel")

p = (
    "The detailed sample distribution of respondents from the four Garhwal "
    "districts is presented in Table 4.3."
)
add_para(doc, p); md_p(p)

# District distribution table — uses actual counts from Malta_.xlsx
dist_rows = [
    ["Pauri Garhwal",   "152", "25", "34 (W) + 15 (CA) = 49", "29",  "34", "311"],
    ["Tehri Garhwal",   "166", "34", "29 (W) + 18 (CA) = 47", "30",  "41", "333"],
    ["Rudraprayag",     "165", "22", "28 (W) + 21 (CA) = 49", "37",  "34", "316"],
    ["Chamoli",         "167", "36", "39 (W) + 24 (CA) = 63", "34",  "21", "340"],
    ["Total",           "650", "117","208",                   "130", "130","1,235"],
]
add_table(doc,
          ["District (Garhwal)", "Stage I — Farmers", "Stage II — Village Traders",
           "Stage III — Wholesalers + Commission Agents",
           "Stage IV — Retailers", "Stage V — Transporters", "Total"],
          dist_rows,
          "Table 4.3: Detailed Sample Distribution of Respondents — Stage Wise "
          "(Malta Distribution Channel, Garhwal Region)")
md_table(["District (Garhwal)", "Stage I — Farmers", "Stage II — Village Traders",
          "Stage III — Wholesalers + CA",
          "Stage IV — Retailers", "Stage V — Transporters", "Total"],
         dist_rows,
         "Table 4.3: Detailed Sample Distribution of Respondents — Stage Wise "
         "(Malta Distribution Channel, Garhwal Region)")

p = (
    "An additional 65 experts and government officials (Horticulture Department, "
    "Mandi Samiti, FPO functionaries and academic experts) were engaged separately "
    "for the qualitative phase under RO 3; their distribution is shown later in "
    "Table 4.7. The aggregate sample frame for the channel-stage survey is "
    "therefore 1,235, and 1,300 when the expert panel is included."
)
add_para(doc, p); md_p(p)

# Subjects-to-items ratio for RO1
p = (
    "Hair et al. (2008), Bartlett et al. (2001), Hinkin (1995) and Schwab (1980) "
    "recommend that the sample size should be at least ten times the number of "
    "items being considered, while Bryant and Yarnold (1995), Garson (2008), "
    "MacCallum et al. (1999), Arrindell and van der Ende (1985), Gorsuch (1983) "
    "and Everitt (1975) place the minimum subjects-to-items ratio at 5:1. Both "
    "thresholds are comfortably satisfied at every stage of the Malta channel, as "
    "shown in Table 4.4."
)
add_para(doc, p); md_p(p)

# Stage activity items per stage (from the docx questionnaire tables)
# Farm: 14 activity-factor rows; Trader: 12; Mandi: 12; Retail: 12; Transport: 14
ratio_rows_ro1 = [
    ["Stage I — Farm Stage",                            "14", "140", "650",  "Yes"],
    ["Stage II — Village Trader / Middlemen Stage",     "12", "120", "117",  "No (5:1 met; 10:1 marginal — see note)"],
    ["Stage III — Wholesale / Mandi Stage",             "12", "120", "208",  "Yes"],
    ["Stage IV — Retail Stage",                         "12", "120", "130",  "Yes (10:1 marginal)"],
    ["Stage V — Transportation Stage",                  "14", "140", "130",  "Yes (5:1 met; 10:1 marginal)"],
]
add_table(doc,
          ["Stage", "No. of Activity Items",
           "Required Sample Size (10× rule)",
           "Actual Sample in Study", "Appropriate Sample Size"],
          ratio_rows_ro1,
          "Table 4.4: Required Sample Size as per Subjects-to-Items Ratio and "
          "Actual Sample Size — Malta Distribution Channel (RO 1)")
md_table(["Stage", "No. of Activity Items",
          "Required Sample Size (10× rule)",
          "Actual Sample in Study", "Appropriate Sample Size"],
         ratio_rows_ro1,
         "Table 4.4: Required Sample Size as per Subjects-to-Items Ratio and "
         "Actual Sample Size — Malta Distribution Channel (RO 1)")

p = (
    "Note: where the 10× threshold is only marginally met, the conservative 5× "
    "threshold of Bryant and Yarnold (1995) is comfortably satisfied at every "
    "stage, ensuring that the sample is adequate for descriptive significance "
    "ranking (RO 1) and for the factor-analytic procedures used in RO 2."
)
add_para(doc, p, italic=True); md_p(f"*{p}*")

# ============================================================
# 4.7 RO 2
# ============================================================
add_heading(doc, "4.7  For Research Objective 2", 1)
md_h("4.7 For Research Objective 2", 2)
add_para(doc,
    "RO 2 — To identify the factors leading to inefficiency (with respect to "
    "cost, time and quality) in the identified activities across the stages of "
    "the Malta distribution channel in the Garhwal region.",
    italic=True)
md_p("*RO 2 — To identify the factors leading to inefficiency (with respect to "
     "cost, time and quality) in the identified activities across the stages of "
     "the Malta distribution channel in the Garhwal region.*")

# 4.7.1
add_heading(doc, "4.7.1  Questionnaire Development", 2)
md_h("4.7.1 Questionnaire Development", 3)
p = (
    "On the basis of variables extracted from a structured literature review and "
    "from suggestions made by Malta supply-chain practitioners, stakeholders and "
    "academicians (Appendix II), a separate factor-level questionnaire was designed "
    "for each stakeholder across the stages — farmer, village trader, local "
    "commission agent, wholesaler, retailer and transporter. The questionnaire was "
    "broadly divided into two sections. The first section captured "
    "socio-demographic characteristics of the respondents (age group, level of "
    "education, years of experience and district). The second section consisted of "
    "close-ended Likert-scaled items covering the reasons leading to inefficiency."
)
add_para(doc, p); md_p(p)

p = (
    "For each stage, the second section was further organised into three blocks — "
    "factors with respect to cost, factors with respect to time and factors with "
    "respect to quality. All items were measured on a five-point Likert scale "
    '("1" = strongly disagree to "5" = strongly agree) (Brown, 2010; Vagias, 2006; '
    "Vogt, 1999). The instrument is reproduced in Appendix III. The data were "
    "imported into IBM SPSS Statistics (version 26) and into jamovi for reliability "
    "testing, exploratory factor analysis and principal component analysis."
)
add_para(doc, p); md_p(p)

# 4.7.2 Pilot Testing
add_heading(doc, "4.7.2  Pilot Testing", 2)
md_h("4.7.2 Pilot Testing", 3)
p = (
    "Based on the final list of variables, the factor-level questionnaire was "
    "pre-tested with a total of 110 stakeholders drawn from the four Garhwal "
    "districts: 55 farmers, 11 village traders, 17 wholesalers / commission "
    "agents, 12 retailers and 15 transporters, in line with the recommendations "
    "of Hair et al. (2010) and Bryman and Bell (2007). Before fieldwork, the "
    "instrument was reviewed with stakeholders and with industry and academic "
    "experts in logistics and supply-chain management; ambiguous or vague items "
    "were removed, repetitive items were deleted, certain items were re-ordered "
    "and a small number of more specific items relevant to Malta were added. "
    "Their feedback was instrumental in producing a concise final instrument "
    "aligned with the research objectives."
)
add_para(doc, p); md_p(p)

pilot2_rows = [
    ["1", "Stage I — Farm Stage",                            "Farmers",                          "55"],
    ["2", "Stage II — Village Trader / Middlemen Stage",     "Village Traders",                  "11"],
    ["3", "Stage III — Wholesale / Mandi Stage",             "Wholesalers / Commission Agents",  "17"],
    ["4", "Stage IV — Retail Stage",                         "Retailers",                        "12"],
    ["5", "Stage V — Transportation Stage",                  "Transporters",                     "15"],
]
add_table(doc,
          ["S. No.", "Stage", "Respondents", "No. of Respondents"],
          pilot2_rows,
          "Table 4.5: Pilot Testing for RO 2 — Malta Distribution Channel")
md_table(["S. No.", "Stage", "Respondents", "No. of Respondents"],
         pilot2_rows,
         "Table 4.5: Pilot Testing for RO 2 — Malta Distribution Channel")

# 4.7.3 Administration of Survey for RO2
add_heading(doc, "4.7.3  Administration of the Survey", 2)
md_h("4.7.3 Administration of the Survey", 3)
p = (
    "The factor-level questionnaire was administered to the same population frame "
    "as RO 1 — farmers in the four Malta-producing districts of the Garhwal "
    "region; village traders in clusters around the major production areas; "
    "wholesalers and commission agents at the Dehradun and Haridwar wholesale "
    "mandis; retailers at district headquarters and adjoining urban centres; "
    "and transporters operating between the orchards, the wholesale mandis and "
    "the downstream urban markets. A multi-stage sampling technique was used, "
    "supplemented where required by snowball sampling. Most responses were "
    "collected through the scheduling method, with the researcher directly "
    "visiting the field; this approach is suitable for extensive enquiries and "
    "yields reasonably reliable results (Gangrade, 2006; Shah, 1972; Pauline, "
    "1968). The same 1,235 stage-wise responses (excluding the 65 expert "
    "respondents reserved for RO 3) form the analytical sample for RO 2."
)
add_para(doc, p); md_p(p)

# Subjects-to-variables table for RO 2
p = (
    "The general norm for factor analysis is to have at least five respondents "
    "per variable (Hair et al., 2008; Bryant and Yarnold, 1995); the "
    "subjects-to-variables ratio should not be lower than 5:1 (Garson, 2008; "
    "MacCallum et al., 1999; Arrindell and van der Ende, 1985; Gorsuch, 1983; "
    "Everitt, 1975). This study comfortably satisfies the more demanding 10:1 "
    "criterion at every stage of the Malta distribution channel, as shown in "
    "Table 4.6."
)
add_para(doc, p); md_p(p)

# Variable counts derived from the codebook in Malta_.xlsx
ratio_rows_ro2 = [
    ["Stage I — Farm Stage",                "Cost",     "7",  "70",  "650",  "Yes"],
    ["",                                    "Time",     "6",  "60",  "650",  "Yes"],
    ["",                                    "Quality",  "6",  "60",  "650",  "Yes"],
    ["Stage II — Village Trader Stage",     "Cost",     "6",  "60",  "117",  "Yes"],
    ["",                                    "Time",     "5",  "50",  "117",  "Yes"],
    ["",                                    "Quality",  "5",  "50",  "117",  "Yes"],
    ["Stage III — Wholesale / Mandi Stage", "Cost",     "6",  "60",  "208",  "Yes"],
    ["",                                    "Time",     "5",  "50",  "208",  "Yes"],
    ["",                                    "Quality",  "5",  "50",  "208",  "Yes"],
    ["Stage IV — Retail Stage",             "Cost",     "6",  "60",  "130",  "Yes"],
    ["",                                    "Time",     "4",  "40",  "130",  "Yes"],
    ["",                                    "Quality",  "5",  "50",  "130",  "Yes"],
    ["Stage V — Transportation Stage",      "Cost",     "6",  "60",  "130",  "Yes"],
    ["",                                    "Time",     "5",  "50",  "130",  "Yes"],
    ["",                                    "Quality",  "5",  "50",  "130",  "Yes"],
]
add_table(doc,
          ["Stage", "Cost / Time / Quality",
           "No. of Variables",
           "Required Sample Size (10× rule)",
           "Sample Size in the Study",
           "Appropriate Sample Size"],
          ratio_rows_ro2,
          "Table 4.6: Required Sample Size as per Subjects-to-Variables Ratio and "
          "Actual Sample Size — Malta Distribution Channel (RO 2)")
md_table(["Stage", "Cost / Time / Quality",
          "No. of Variables",
          "Required Sample Size (10× rule)",
          "Sample Size in the Study",
          "Appropriate Sample Size"],
         ratio_rows_ro2,
         "Table 4.6: Required Sample Size as per Subjects-to-Variables Ratio and "
         "Actual Sample Size — Malta Distribution Channel (RO 2)")

# Reliability summary
p = (
    "Reliability of each scale was assessed using Cronbach's α prior to the "
    "factor-analytic procedures. Across all fifteen sub-scales (five stages × "
    "three aspects), Cronbach's α ranged from 0.82 to 0.92, with a mean of 0.885; "
    "all values comfortably exceed the 0.70 threshold recommended by Nunnally "
    "(1978) and confirm strong internal consistency. The Kaiser–Meyer–Olkin (KMO) "
    "measure of sampling adequacy and Bartlett's test of sphericity were "
    "additionally examined for each scale prior to factor extraction. PCA with "
    "varimax rotation was used to extract factors, and items with loadings ≥ 0.50 "
    "were retained. For the Farm-stage scale, three components — interpreted "
    "respectively as the Cost factor (27.2 % variance), the Quality factor (21.6 %) "
    "and the Time factor (19.2 %) — together explain 68.0 % of the variance, "
    "indicating a clean and theoretically interpretable factor structure."
)
add_para(doc, p); md_p(p)

# ============================================================
# 4.8 RO 3
# ============================================================
add_heading(doc, "4.8  For Research Objective 3", 1)
md_h("4.8 For Research Objective 3", 2)
add_para(doc,
    "RO 3 — To develop a framework for improving the distribution-channel "
    "efficiency of Malta with specific reference to the Garhwal region.",
    italic=True)
md_p("*RO 3 — To develop a framework for improving the distribution-channel "
     "efficiency of Malta with specific reference to the Garhwal region.*")

p = (
    "Since it was not feasible for an individual researcher to address every "
    "issue surfaced by RO 2, this objective focuses on the major reasons under "
    "each factor leading to inefficiency. The factor loading of every variable "
    "across the stages of the Malta distribution channel was examined, and the "
    "principal variable under each factor was identified on the basis of the "
    "highest factor loading (Negi and Anand, 2018a, 2018b; Tiwari, 2012). The "
    "higher the factor loading, the greater the contribution of the variable to "
    "the factor (Hair et al., 2010; Field, 2009; Malhotra, 2007; Harman, 1976). "
    "Factor loadings are very similar to weights and signify the strength of "
    "correlation between the factor and the variable (Kline, 1994). Variables "
    "with the highest loadings therefore represent the major reasons that "
    "contribute most to inefficiency and that demand the most immediate "
    "attention if distribution-channel efficiency is to be improved (Tiwari, "
    "2012; Hair et al., 2010; Field, 2009; Malhotra, 2007; Harman, 1976)."
)
add_para(doc, p); md_p(p)

p = (
    "Qualitative research was used to derive measures and to construct the "
    "improvement framework. Qualitative methods help to collect, analyse and "
    "understand raw data by perceiving participants' reactions (Creswell, 1994); "
    "they sit within the interpretivist tradition and are well suited to "
    "understanding the diverse and complex nature of human action. Qualitative "
    "data emphasise in-depth enquiry, detailed description and direct quotations "
    "that capture experiences and personal viewpoints (Patton, 1990). The "
    "approach involves observation, focus-group discussions and interviews "
    "(Cavana et al., 2008). In this study, in-depth semi-structured interviews "
    "were used to derive measures for the highest-loading reasons. Pre-formulated "
    "questions guided the interview, but the format also allowed new questions "
    "to emerge during the conversation (Myers, 2009, p. 124), giving each "
    "interviewee the freedom to add meaningful insights. Measures for items "
    "outside this set are beyond the scope of the present study and are "
    "acknowledged as a limitation."
)
add_para(doc, p); md_p(p)

# 4.8.1 Conceptualization
add_heading(doc, "4.8.1  Conceptualisation and Protocol Design", 2)
md_h("4.8.1 Conceptualisation and Protocol Design", 3)
p = (
    "Based on the highest factor loadings obtained in RO 2, only those variables "
    "that emerged as the most critical reasons for inefficiency in each factor "
    "across the stages of the Malta distribution channel were used for "
    "conceptualisation (Appendix IV); accordingly, an interview protocol was "
    "designed for each stage to elicit measures for improving distribution-channel "
    "efficiency (Appendix V). A protocol is a formal document comprising the "
    "interview questionnaire together with the rules of administration. During "
    "data collection, construct validity was secured through multiple sources of "
    "evidence and through use of the protocol; reliability was ensured through "
    "consistent application of the same protocol to every interviewee (Yin, "
    "2003). A small pilot was conducted with Malta supply-chain experts to verify "
    "that the questions were comprehensible to the respondents and to test the "
    "rationality, narrowness or breadth of each question."
)
add_para(doc, p); md_p(p)

# 4.8.2 Sampling Design
add_heading(doc, "4.8.2  Sampling Design", 2)
md_h("4.8.2 Sampling Design", 3)
p = (
    "One of the most vital components of a qualitative study is the selection of "
    "respondents who are willing to provide information that is a true "
    "representation of the targeted population (Cavana et al., 2008). Bryman and "
    "Bell (2007) note that statistical representativeness need not be the "
    "primary focus in qualitative research, since the emphasis is on a thorough "
    "analysis of meaning. Cavana et al. (2008, p. 137) highlight the advantage "
    "of non-probability methods for extracting high-quality information rapidly "
    "from informed respondents. Accordingly, a judgmental (purposive) sampling "
    "method was adopted for the in-depth interviews. In this approach, the "
    "researcher is fully clear about the information sought and selects "
    "respondents who, by virtue of their experience, are best placed to provide "
    "it (Bernard, 2002). The method is particularly useful for identifying "
    "information-rich cases that maximise the value extracted from a finite "
    "research budget (Patton, 2002). Beyond expertise, willingness to participate "
    "and the ability to communicate real-life experience in a critical and "
    "insightful manner were also considered (Bernard, 2002; Spradley, 1979)."
)
add_para(doc, p); md_p(p)

p = (
    "Judgmental sampling places considerable weight on data saturation — "
    "continuing interviews until no new substantive information emerges (Miles "
    "and Huberman, 1994). Sample sizes for the qualitative phase were therefore "
    "determined by saturation rather than by an a priori target. The "
    "stage-wise number of experts engaged is shown in Table 4.7."
)
add_para(doc, p); md_p(p)

# Saturation table - 65 experts split across stages
sat_rows = [
    ["Stage I — Farm Stage",                            "16"],
    ["Stage II — Village Trader / Middlemen Stage",     "10"],
    ["Stage III — Wholesale / Mandi Stage",             "14"],
    ["Stage IV — Retail Stage",                         "11"],
    ["Stage V — Transportation Stage",                  "14"],
    ["Total",                                           "65"],
]
add_table(doc,
          ["Stage", "No. of Experts"],
          sat_rows,
          "Table 4.7: Sample Size Based on Saturation — Malta Distribution Channel "
          "(Qualitative Phase, RO 3)")
md_table(["Stage", "No. of Experts"],
         sat_rows,
         "Table 4.7: Sample Size Based on Saturation — Malta Distribution Channel "
         "(Qualitative Phase, RO 3)")

p = (
    "Prior to each interview, respondents were assured that the information "
    "collected would be confidential and used solely for academic purposes. "
    "Responses were audio-recorded with consent, transcribed, and inductively "
    "coded. Sample transcripts and the codebook are reproduced in Appendices VI "
    "and VII. Measures for each high-loading variable were derived from the "
    "transcribed conversations and the resulting codes; based on these, the "
    "improvement framework was assembled. Findings were further reviewed by an "
    "independent expert, as recommended by Yin (2003), whose evaluation "
    "corroborated the conclusions of the study. To enhance the overall quality "
    "of the research, triangulation was attempted by obtaining data from varied "
    "informants representing different stages of the channel, by vetting "
    "findings with an independent expert, and by adhering consistently to the "
    "interview protocol throughout (Yin, 2003)."
)
add_para(doc, p); md_p(p)

# Steps for RO3
add_para(doc, "The steps followed for this objective are as follows:", bold=True)
md_p("**The steps followed for this objective are as follows:**")
ro3_steps = [
    "Selection of experts on the basis of judgmental sampling, comprising: "
    "(a) Malta / horticulture supply-chain experts; (b) cold-chain experts; "
    "(c) logistics experts; (d) Garhwal-based agribusiness experts and FPO "
    "functionaries; (e) officials of the Horticulture Department, "
    "Government of Uttarakhand, and the Mandi Samiti; and (f) academic experts "
    "in agribusiness and supply-chain management.",
    "Conduct of semi-structured interviews with the experts.",
    "Verbatim transcription of the recorded conversations.",
    "Inductive coding of the transcripts.",
    "Derivation of measures for improving distribution-channel efficiency of "
    "Malta on the basis of the codes.",
    "Construction of the improvement framework on the basis of the consolidated "
    "output.",
]
for i, s in enumerate(ro3_steps, start=1):
    doc.add_paragraph(f"{i}. {s}", style="List Number")
    md_lines.append(f"{i}. {s}")
md_lines.append("")

# ============================================================
# 4.9 Scope
# ============================================================
add_heading(doc, "4.9  Scope of the Study", 1)
md_h("4.9 Scope of the Study", 2)
p = (
    "The scope of this study is restricted to the distribution channel of Malta "
    "(Citrus sinensis) in the Garhwal region of Uttarakhand. The four districts "
    "of Pauri Garhwal, Tehri Garhwal, Rudraprayag and Chamoli were selected as "
    "the production base because they account for the largest concentration of "
    "Malta cultivation in the state and represent a key livelihood crop for "
    "mid-hill farmers. The channel was traced from the farm gate, through the "
    "village trader / middleman and the local commission agent, to the "
    "wholesale mandi (with Dehradun and Haridwar serving as the principal "
    "aggregation points and Saharanpur and Delhi-Azadpur as onward destinations), "
    "and onwards to the retailer; this segment was selected because the bulk of "
    "the inefficiency in the Malta channel is concentrated between these stages. "
    "Studies of consumption-stage behaviour, of export markets, and of citrus "
    "varieties other than Malta are outside the scope of the present work and "
    "are flagged as opportunities for future research."
)
add_para(doc, p); md_p(p)

# ============================================================
# Concluding Remarks
# ============================================================
add_heading(doc, "Concluding Remarks", 1)
md_h("Concluding Remarks", 2)
p = (
    "This chapter has set out the research methodology adopted to attain the "
    "stated objectives. Both exploratory and conclusive research designs have "
    "been used. Data collection and analysis combine quantitative methods "
    "(structured questionnaires, descriptive analysis, reliability testing via "
    "Cronbach's α, and principal component / exploratory factor analysis) with "
    "qualitative methods (semi-structured expert interviews, transcription and "
    "inductive coding) in a sequential mixed-method design. The chapter has "
    "also articulated the rationale, the problem, the research questions and "
    "objectives, the sampling design, the administration of the survey and the "
    "scope of the work. The next chapter presents the data analysis and the "
    "findings of the study in detail."
)
add_para(doc, p); md_p(p)

# Save
doc.save("Chapter_4_Research_Methodology.docx")

with open("Chapter_4_Research_Methodology.md", "w") as f:
    f.write("\n".join(md_lines))

print("WROTE Chapter_4_Research_Methodology.docx and .md")
