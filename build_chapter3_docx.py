"""
Build Chapter_3_Research_Methodology.docx from the same content as the Markdown file.
Uses python-docx for proper headings, justified body text, grid-style tables,
and a formatted reference list.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --------------------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------------------
def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def shade_cell(cell, color_hex="D9E1F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_table(doc, headers, rows, col_widths=None, header_shade="305496",
              header_font_color=RGBColor(0xFF, 0xFF, 0xFF), font_size=10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = header_font_color
        run.font.name = "Times New Roman"
        shade_cell(hdr_cells[i], header_shade)
        set_cell_borders(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = "Times New Roman"
            # Numeric cells right-aligned, others left
            try:
                float(str(val).replace(",", "").replace("%", ""))
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except ValueError:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Bold totals
            if "Total" in str(val) or "Cumulative" in str(val) or "Pooled" in str(val):
                run.bold = True
            set_cell_borders(cells[c_idx])
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Shade total rows
        if any(("Total" in str(v) or "Cumulative" in str(v) or "Pooled" in str(v))
               for v in row):
            for c in cells:
                shade_cell(c, "F2F2F2")
    return table


def style_paragraph(p, size=12, font="Times New Roman", justify=True,
                    space_after=6, line_spacing=1.5):
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    for r in p.runs:
        r.font.name = font
        r.font.size = Pt(size)


def add_para(doc, text, **kwargs):
    p = doc.add_paragraph(text)
    style_paragraph(p, **kwargs)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        if level == 0:
            r.font.size = Pt(20)
        elif level == 1:
            r.font.size = Pt(16)
        elif level == 2:
            r.font.size = Pt(13)
        else:
            r.font.size = Pt(12)
    return h


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_source(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    return p


def add_reference(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-1.0)  # hanging indent
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
    return p


# --------------------------------------------------------------------------------------
# Build document
# --------------------------------------------------------------------------------------
doc = Document()

# Set default font and margins
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3.0)
    s.right_margin = Cm(2.5)

# ---- Title ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("CHAPTER 3")
r.bold = True
r.font.size = Pt(22)
r.font.name = "Times New Roman"
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("RESEARCH METHODOLOGY")
r.bold = True
r.font.size = Pt(20)
r.font.name = "Times New Roman"
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
subtitle.paragraph_format.space_after = Pt(24)

# ============================================================================
# 3.1 Introduction
# ============================================================================
add_heading(doc, "3.1 Introduction", level=1)
add_para(doc,
    "Research methodology is the systematic blueprint that translates a research problem into "
    "investigable questions and converts those questions into empirical evidence (Kothari, 2004; "
    "Creswell & Creswell, 2018). It is, in Saunders, Lewis, and Thornhill's (2019) widely cited "
    "\"research onion\" formulation, the layered choice set comprising philosophy, approach, "
    "strategy, choices, time horizon, and techniques that together justify how a study generates "
    "defensible knowledge.")
add_para(doc,
    "This chapter operationalises the methodological choices for the present investigation "
    "entitled \"A Study on Factors Impacting Optimisation of Distribution Channels of Malta "
    "(Citrus sinensis) with Special Reference to Garhwal Region.\" The chapter is organised in "
    "alignment with the conceptual flow presented in Figure 3.1: (i) clarification of the research "
    "design; (ii) description of the study area; (iii) sampling design; (iv) data-collection "
    "instruments and procedure; (v) the funnelling of an exhaustive factor inventory of 144 items "
    "down to a parsimonious, statistically validated set of 82 items; (vi) the analytical framework, "
    "which integrates Fuzzy Analytic Hierarchy Process (Fuzzy AHP) for criteria weighting and "
    "Fuzzy ELECTRE for ranking distribution-channel alternatives; and (vii) considerations of "
    "validity, reliability, ethics, and methodological limitations.")
add_para(doc,
    "The methodology is deliberately pluralistic but converging. Quantitative survey evidence from "
    "supply-chain stakeholders is triangulated with qualitative semi-structured expert interviews "
    "(Patton, 2015; Yin, 2018). The use of fuzzy multi-criteria decision-making (F-MCDM) tools is "
    "motivated by two characteristic features of the Malta supply chain in Garhwal: (a) human "
    "judgements expressed in linguistic, imprecise terms (\"high cost\", \"moderate delay\") that "
    "are poorly captured by crisp numbers (Zadeh, 1965; Bellman & Zadeh, 1970); and (b) the "
    "simultaneous presence of multiple, often conflicting, evaluation criteria - cost, time, and "
    "quality - across five sequential supply-chain stages (Chang, 1996; Hatami-Marbini & Tavana, 2011).")

# ============================================================================
# 3.2 Research Design
# ============================================================================
add_heading(doc, "3.2 Research Design", level=1)

add_heading(doc, "3.2.1 Research Philosophy and Approach", level=2)
add_para(doc,
    "The study adopts a pragmatist philosophical stance (Tashakkori & Teddlie, 2010), which "
    "permits the researcher to use whichever method - quantitative, qualitative, or both - best "
    "answers the research question. Epistemologically, the work is post-positivist for the "
    "structured survey (objective measurement of stakeholder perceptions) and interpretivist for "
    "the expert-interview component (meaning-making around supply-chain inefficiencies). The "
    "reasoning is abductive: factors are first inferred from theory and literature, then refined "
    "through field engagement, and finally tested through structured measurement (Saunders et al., 2019).")

add_heading(doc, "3.2.2 Type of Research Design", level=2)
add_para(doc, "A cross-sectional, descriptive-cum-analytical research design is employed.")
for bullet in [
    "Descriptive - because the study profiles stakeholders, maps the existing Malta distribution "
    "channels in Garhwal, and documents the prevailing inefficiencies (Malhotra & Dash, 2016).",
    "Analytical - because it goes beyond description to quantify and rank the relative impact of "
    "factors on channel optimisation, using Fuzzy AHP and Fuzzy ELECTRE (Govindan et al., 2015).",
    "Cross-sectional - because data were collected from each respondent at a single point in time "
    "during one Malta marketing season (October 2023 - February 2024), rather than longitudinally "
    "(Bryman & Bell, 2015). Cross-sectional designs are appropriate when the objective is to "
    "identify and prioritise determinants rather than track change over time (Levin, 2006)."
]:
    p = doc.add_paragraph(bullet, style="List Bullet")
    style_paragraph(p, space_after=4)

add_heading(doc, "3.2.3 Justification of the Mixed-Methods, Cross-Sectional Choice", level=2)
add_caption(doc, "Table 3.0  Justification of Methodological Choices")
add_table(doc,
    headers=["Methodological Issue", "Justification"],
    rows=[
        ["Why cross-sectional?",
         "Malta is a seasonal crop; capturing one full marketing cycle yields complete "
         "supply-chain evidence within feasible time and budget (Sedgwick, 2014)."],
        ["Why mixed-methods?",
         "Quantitative ranking alone cannot explain why factors operate; qualitative expert "
         "insight cannot generalise. Sequential triangulation overcomes both gaps "
         "(Creswell & Plano Clark, 2017)."],
        ["Why fuzzy MCDM?",
         "Stakeholder judgements about supply-chain inefficiency are inherently linguistic and "
         "uncertain; fuzzy logic preserves this nuance (Zadeh, 1965; Kahraman et al., 2003)."],
        ["Why two MCDM tools (AHP + ELECTRE)?",
         "AHP yields weights; ELECTRE yields outranking-based rankings. Combining them generates "
         "both criterion importance and alternative prioritisation (Roy, 1991; "
         "Hatami-Marbini & Tavana, 2011)."],
    ],
    col_widths=[Cm(5.5), Cm(11.0)])
add_source(doc, "Source: Author's compilation.")

add_para(doc,
    "Figure 3.1 (provided in the thesis) presents the methodological flow visually: Preliminary "
    "Literature -> Problem Definition & Theoretical Framework -> Selection of Methods (Survey + "
    "Interview) -> Parallel pipelines for survey-based factor ranking (sampling -> questionnaire "
    "-> pilot -> administration -> editing -> ranking) and expert interviews (purposive sampling "
    "-> semi-structured guide -> Fuzzy ELECTRE / transcribed coding) -> Convergence on factors "
    "driving distribution-channel inefficiency -> Interpretation -> Conclusions and recommendations.")

# ============================================================================
# 3.3 Study Area Profile
# ============================================================================
add_heading(doc, "3.3 Study Area Profile", level=1)

add_heading(doc, "3.3.1 Geographical Characteristics of the Garhwal Region", level=2)
add_para(doc,
    "The Garhwal region constitutes the western half of the state of Uttarakhand in the Indian "
    "Himalayas, lying approximately between 29°26'-31°28' N latitude and 77°34'-80°06' E longitude "
    "(Government of Uttarakhand, 2022). It encompasses seven administrative districts; this study "
    "focuses on the four districts in which Malta is cultivated commercially: Pauri Garhwal, Tehri "
    "Garhwal, Chamoli, and Rudraprayag. Elevation ranges from ~600 m to over 3,000 m AMSL. Malta "
    "thrives in the mid-hill belt of 800-1,800 m, where temperatures range from 10 °C (winter) to "
    "32 °C (summer) and annual rainfall is 1,000-2,200 mm (ICAR-CITH, 2021).")

add_heading(doc, "3.3.2 Demographic Profile", level=2)
add_para(doc,
    "According to the Census of India (2011) projected to 2023, the four study districts together "
    "house ~3.2 million people, of whom over 70% reside in rural areas and depend on agriculture "
    "and horticulture for livelihood (Directorate of Economics & Statistics, Uttarakhand, 2022).")

add_heading(doc, "3.3.3 Agricultural and Horticultural Profile", level=2)
add_para(doc,
    "Uttarakhand contributes approximately 38,000 metric tonnes of Malta annually, with Garhwal "
    "accounting for ~62% of state production (NHB, 2023). The crop is grown predominantly on small "
    "and marginal holdings (<1 ha), using rain-fed cultivation with limited mechanisation "
    "(Pant et al., 2020).")

add_heading(doc, "3.3.4 Malta (Citrus sinensis) Cultivation in the Region", level=2)
add_para(doc,
    "Malta was introduced to Garhwal in the early 20th century and is today a socio-economically "
    "strategic horticultural crop for hill farmers (Kumar et al., 2019). Two principal varieties "
    "dominate: Sweet Malta and Jaffa Malta. The fruit is harvested between November and February, "
    "has a short shelf-life of 7-14 days under ambient conditions, and depends on a fragile, "
    "multi-tier distribution channel (Singh & Bhatt, 2018).")

add_heading(doc, "3.3.5 Marketing Infrastructure", level=2)
add_para(doc,
    "The principal regulated mandis serving Malta are at Srinagar (Pauri), New Tehri, Gopeshwar, "
    "and Rudraprayag, supplemented by smaller collection centres. Channel members include farmers, "
    "village traders, commission agents, wholesalers, retailers, and transporters. Cold-storage "
    "capacity for citrus in Garhwal is reported to be less than 5% of seasonal arrivals "
    "(HAPPRC, 2021), which underscores the practical importance of optimising the distribution channel.")

add_caption(doc, "Table 3.1  Selected Profile of the Four Study Districts")
add_table(doc,
    headers=["District", "Geographical Area (km²)", "Mid-hill Malta belt",
             "Principal Mandi", "Estimated Malta Output (MT, 2022-23)"],
    rows=[
        ["Pauri Garhwal", "5,329", "900-1,700 m", "Srinagar", "~7,800"],
        ["Tehri Garhwal", "3,642", "1,000-1,800 m", "New Tehri", "~6,500"],
        ["Chamoli", "8,030", "1,100-1,800 m", "Gopeshwar", "~5,400"],
        ["Rudraprayag", "1,984", "1,000-1,700 m", "Rudraprayag", "~3,900"],
    ],
    col_widths=[Cm(3.0), Cm(3.5), Cm(3.5), Cm(3.0), Cm(3.5)])
add_source(doc, "Source: Compiled from NHB (2023); Directorate of Horticulture, Uttarakhand (2022); HAPPRC (2021).")

# ============================================================================
# 3.4 Sampling Design
# ============================================================================
add_heading(doc, "3.4 Sampling Design", level=1)

add_heading(doc, "3.4.1 Population Definition", level=2)
add_para(doc,
    "The target population comprises all individuals participating in the Malta distribution "
    "channel in the four study districts of Garhwal, namely: (i) Malta growers/farmers; (ii) "
    "village traders / pre-harvest contractors; (iii) commission agents (arhatiyas); (iv) "
    "wholesalers operating in regulated mandis; (v) retailers (mandi-tied and non-mandi); "
    "(vi) transporters; and (vii) horticulture officials and subject-matter experts (academic, "
    "KVK, and extension functionaries).")

add_heading(doc, "3.4.2 Sampling Technique", level=2)
add_para(doc,
    "A multi-stage, stratified random sampling design with a complementary purposive component "
    "for experts was used (Cochran, 1977; Etikan et al., 2016).")
for bullet in [
    "Stage 1 - District selection (purposive): the four Garhwal districts with significant Malta cultivation.",
    "Stage 2 - Block/Village selection (stratified random): within each district, blocks were "
    "stratified by elevation belt and three blocks per district were drawn at random. Two "
    "Malta-growing villages per block were then selected probability-proportional-to-size (PPS) "
    "on the basis of the number of Malta-cultivating households.",
    "Stage 3 - Respondent selection within each village/mandi (simple random): farmers from "
    "village rosters; traders, agents, wholesalers, retailers, and transporters from mandi "
    "registration lists.",
    "Expert sub-sample (purposive): selected on the basis of (a) >= 10 years' Malta-related "
    "experience, (b) institutional affiliation (KVK, Directorate of Horticulture, GBPUA&T, "
    "HAPPRC), and (c) willingness to participate (Patton, 2015).",
]:
    p = doc.add_paragraph(bullet, style="List Bullet")
    style_paragraph(p, space_after=4)

add_heading(doc, "3.4.3 Sample-Size Determination", level=2)
add_para(doc,
    "For the survey strand, sample size was computed using Cochran's (1977) formula for a "
    "large/unknown population:")
formula = doc.add_paragraph()
formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = formula.add_run("n0 = (z² × p × q) / e²")
r.italic = True
r.font.size = Pt(13)
r.font.name = "Cambria Math"
add_para(doc,
    "with z = 1.96 (95% CI), p = 0.5 (most-conservative variance), q = 1 - p = 0.5, and "
    "e = 0.04 (4% margin), yielding n0 ≈ 600. Because seven heterogeneous stakeholder strata were "
    "sampled, the figure was inflated to n = 1,300 to ensure adequate cell sizes for stratum-wise "
    "reliability and PCA, and to satisfy the \"10 respondents per measured item\" rule of thumb "
    "for factor analysis (Hair et al., 2019). The achieved sample of N = 1,300 valid responses "
    "therefore exceeds the minimum required for both inferential statistics and the application "
    "of Fuzzy AHP/Fuzzy ELECTRE.")

add_heading(doc, "3.4.4 Selection of Districts/Blocks/Villages", level=2)
add_para(doc,
    "Twelve blocks (3 per district), 24 villages, and four major mandis were finally surveyed. "
    "Expert interviews were conducted across all four districts plus an additional set with "
    "state-level officials and academic experts.")

add_heading(doc, "3.4.5 Achieved Sample Distribution", level=2)
add_caption(doc, "Table 3.2  Distribution of Respondents by Stakeholder Category and District (N = 1,300)")
add_table(doc,
    headers=["Stakeholder Category", "Pauri Garhwal", "Tehri Garhwal", "Chamoli",
             "Rudraprayag", "Total"],
    rows=[
        ["Farmers / Growers", "152", "166", "167", "165", "650"],
        ["Village Traders", "25", "34", "36", "22", "117"],
        ["Commission Agents", "15", "18", "24", "21", "78"],
        ["Wholesalers (Mandi)", "34", "29", "39", "28", "130"],
        ["Retailers", "29", "30", "34", "37", "130"],
        ["Transporters", "34", "41", "21", "34", "130"],
        ["Experts / Officials", "22", "15", "19", "9", "65"],
        ["Total", "311", "333", "340", "316", "1,300"],
    ],
    col_widths=[Cm(4.0), Cm(2.6), Cm(2.6), Cm(2.4), Cm(2.6), Cm(2.0)])
add_source(doc, "Source: Field survey (2023-24).")
add_para(doc,
    "The dominant share of farmers (50%) reflects their numerical preponderance in the Malta value "
    "chain, while the remaining 50% covers the downstream channel members and experts whose "
    "perceptions are essential for an end-to-end optimisation analysis.")

# ============================================================================
# 3.5 Data Collection
# ============================================================================
add_heading(doc, "3.5 Data Collection", level=1)

add_heading(doc, "3.5.1 Sources of Data", level=2)
for bullet in [
    "Primary data: structured questionnaire administered to 1,235 channel-member respondents and "
    "semi-structured interview schedule administered to 65 experts/officials.",
    "Secondary data: publications of the National Horticulture Board (NHB), the Directorate of "
    "Horticulture (Uttarakhand), ICAR-CITH (Lucknow), HAPPRC (Srinagar Garhwal), the Census of "
    "India, peer-reviewed journals, and theses retrieved from Shodhganga.",
]:
    p = doc.add_paragraph(bullet, style="List Bullet")
    style_paragraph(p, space_after=4)

add_heading(doc, "3.5.2 Survey Instrument Design", level=2)
add_para(doc,
    "The questionnaire (see Appendix-I, Malta-Distribution-Survey, in the project repository) was "
    "developed in five iterative steps:")
for i, step in enumerate([
    "Item pool generation from a systematic review of supply-chain, agricultural-marketing, and "
    "citrus-distribution literature (e.g., Mentzer et al., 2001; Gunasekaran et al., 2004; "
    "Punjabi, 2009; Negi & Anand, 2015; Chauhan & Singh, 2018).",
    "Field reconnaissance - eight unstructured field visits to Malta orchards, collection centres, "
    "and mandis, used to surface region-specific factors.",
    "Expert pre-screening - six experts (academic + practitioner) reviewed item relevance and "
    "clarity (Lawshe, 1975).",
    "Translation and back-translation between English and Hindi to ensure semantic equivalence "
    "(Brislin, 1970).",
    "Pilot study (see Section 3.7).",
], start=1):
    p = doc.add_paragraph(f"{i}. {step}")
    style_paragraph(p, space_after=4)
add_para(doc, "The instrument is organised in five sections:")
for bullet in [
    "Section A: Respondent profile.",
    "Section B: Stage-specific factor batteries (Farmer, Trader, Wholesaler/Mandi, Retailer, "
    "Transporter), each rated on Cost / Time / Quality dimensions on a 5-point Likert scale "
    "(1 = Not Significant, 5 = Highly Significant).",
    "Section C: Cross-stage coordination and external factors.",
    "Section D: Improvement-priority rankings.",
    "Section E: Variety, seasonality, and shelf-life questions.",
]:
    p = doc.add_paragraph(bullet, style="List Bullet")
    style_paragraph(p, space_after=4)

add_heading(doc, "3.5.3 Interview Schedule for Experts", level=2)
add_para(doc,
    "A semi-structured interview guide with eleven open-ended prompts was used (Kvale & Brinkmann, "
    "2015). Themes covered policy, infrastructure, technology adoption, cooperative organisation, "
    "and recommended channel re-design.")

add_heading(doc, "3.5.4 Field Procedure", level=2)
add_para(doc,
    "Data were collected face-to-face by the researcher and two trained field assistants between "
    "October 2023 and February 2024. Each interview lasted 25-40 minutes (survey) or 60-90 minutes "
    "(expert). Responses were entered into a spreadsheet on the same day and quality-checked weekly.")

add_heading(doc, "3.5.5 Data Editing, Coding, and Cleaning", level=2)
add_para(doc,
    "All instruments were edited for completeness; cases with > 10% missing items were discarded "
    "(Hair et al., 2019). Remaining missing values were imputed by series mean within stakeholder "
    "stratum. The cleaned dataset (Malta_.xlsx, sheet Data_for_jamovi) contains 1,300 cases × 82 "
    "factor items plus six demographic variables.")

# ============================================================================
# 3.6 Identification and Refinement of Factors
# ============================================================================
add_heading(doc, "3.6 Identification and Refinement of Factors: From 144 to 82", level=1)
add_para(doc,
    "A defining feature of this study is the systematic distillation of 144 candidate factors into "
    "82 statistically validated factors. The funnelling logic is summarised in Figure 3.2.")

add_heading(doc, "3.6.1 Stage 1 - Generation of the Initial Pool (144 Items)", level=2)
add_para(doc, "The initial inventory was assembled from three converging sources:")
add_caption(doc, "Table 3.A  Sources of the 144-Item Initial Factor Pool")
add_table(doc,
    headers=["Source", "Number of Items Contributed"],
    rows=[
        ["Systematic literature review (47 papers, 2000-2023)", "78"],
        ["Field observations and unstructured grower/trader interactions", "39"],
        ["Pre-screening with six domain experts (modified Delphi round)", "27"],
        ["Total candidate items (with duplicates)", "144"],
    ],
    col_widths=[Cm(11.0), Cm(4.0)])
add_para(doc,
    "The 144-item draft was deliberately over-inclusive in line with DeVellis's (2017) "
    "recommendation that scale development begin with a pool roughly 3-4 times the intended final length.")

add_heading(doc, "3.6.2 Stage 2 - Content Validity and Expert Screening", level=2)
add_para(doc,
    "Each item was rated by 15 experts for relevance, clarity, and contextual fit on a 4-point "
    "scale. The Content Validity Ratio (CVR) was computed for each item (Lawshe, 1975):")
formula2 = doc.add_paragraph()
formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = formula2.add_run("CVR = (n_e - N/2) / (N/2)")
r.italic = True
r.font.size = Pt(13)
r.font.name = "Cambria Math"
add_para(doc,
    "Items with CVR < 0.49 (Lawshe's critical value for N = 15 at α = 0.05) were dropped. This "
    "step removed 38 items (e.g., redundant or context-irrelevant factors), leaving 106 items.")

add_heading(doc, "3.6.3 Stage 3 - Pilot Reliability Screening", level=2)
add_para(doc,
    "The 106-item pilot instrument was administered to 80 respondents (pilot sample, see Section "
    "3.7). Items whose corrected item-total correlation < 0.30 or whose deletion increased "
    "Cronbach's α were dropped (Nunnally & Bernstein, 1994; Hair et al., 2019). This step removed "
    "a further 14 items, leaving 92 items.")

add_heading(doc, "3.6.4 Stage 4 - Exploratory Factor Analysis on the Main Sample", level=2)
add_para(doc,
    "Principal Component Analysis (Varimax rotation) was performed stage-wise. Items with "
    "communalities < 0.40, loadings < 0.50, or substantial cross-loadings (Δ < 0.20) were dropped "
    "(Hair et al., 2019; Tabachnick & Fidell, 2019). Ten further items were removed, producing the "
    "final 82-item instrument.")

add_heading(doc, "3.6.5 Final 82-Factor Architecture", level=2)
add_para(doc,
    "The retained 82 factors are organised as a 5-stage × 3-dimension matrix (Table 3.3). The "
    "structure matches the Codebook sheet of Malta_.xlsx.")
add_caption(doc, "Table 3.3  Final 82-Factor Architecture by Supply-Chain Stage and Dimension")
add_table(doc,
    headers=["Supply-chain Stage", "Cost factors", "Time factors", "Quality factors", "Stage Total"],
    rows=[
        ["Farmer / Grower", "7", "6", "6", "19"],
        ["Local Trader / Middleman", "6", "5", "5", "16"],
        ["Wholesaler / Mandi", "6", "5", "5", "16"],
        ["Retailer", "6", "4", "5", "15"],
        ["Transporter", "6", "5", "5", "16"],
        ["Dimension Total", "31", "25", "26", "82"],
    ],
    col_widths=[Cm(5.0), Cm(2.6), Cm(2.6), Cm(2.8), Cm(2.5)])
add_source(doc, "Source: Author, derived from Malta_.xlsx Codebook.")
add_para(doc,
    "Figure 3.2 - Funnelling Logic of Factor Refinement: Literature + Field + Delphi (144 items) "
    "-> CVR screening (-38) -> 106 -> Item-total / α-if-deleted (-14) -> 92 -> PCA loadings & "
    "cross-loading checks (-10) -> 82 final factors.")

# ============================================================================
# 3.7 Pilot Study, Validity, and Reliability
# ============================================================================
add_heading(doc, "3.7 Pilot Study, Validity, and Reliability", level=1)

add_heading(doc, "3.7.1 Pilot Study", level=2)
add_para(doc,
    "A pilot was conducted on 80 respondents (~6% of the planned sample) drawn from one block in "
    "each district, in line with Connelly's (2008) \"10% rule.\" The pilot served three purposes: "
    "(i) detect ambiguous wording, (ii) refine field-administration logistics, and (iii) compute "
    "preliminary reliability.")

add_heading(doc, "3.7.2 Reliability - Cronbach's Alpha", level=2)
add_para(doc,
    "Internal consistency for each of the 15 sub-scales (5 stages × 3 dimensions) was assessed on "
    "the full sample of 1,300 cases. All α values exceed the conventional 0.70 threshold "
    "(Nunnally & Bernstein, 1994), and most exceed the 0.80 \"excellent\" mark (Hair et al., 2019).")
add_caption(doc, "Table 3.4  Reliability Statistics for the 15 Sub-Scales (N = 1,300)")
add_table(doc,
    headers=["Sub-scale", "No. of items", "Cronbach's α", "Avg. inter-item r"],
    rows=[
        ["Farm - Cost", "7", "0.917", "0.614"],
        ["Farm - Time", "6", "0.897", "0.591"],
        ["Farm - Quality", "6", "0.919", "0.656"],
        ["Trader - Cost", "6", "0.899", "0.597"],
        ["Trader - Time", "5", "0.847", "0.525"],
        ["Trader - Quality", "5", "0.894", "0.628"],
        ["Mandi - Cost", "6", "0.905", "0.615"],
        ["Mandi - Time", "5", "0.861", "0.553"],
        ["Mandi - Quality", "5", "0.886", "0.610"],
        ["Retail - Cost", "6", "0.885", "0.563"],
        ["Retail - Time", "4", "0.820", "0.533"],
        ["Retail - Quality", "5", "0.875", "0.583"],
        ["Transport - Cost", "6", "0.900", "0.601"],
        ["Transport - Time", "5", "0.903", "0.651"],
        ["Transport - Quality", "5", "0.870", "0.572"],
        ["Pooled (82 items)", "82", "≥ 0.82", "—"],
    ],
    col_widths=[Cm(4.5), Cm(2.5), Cm(3.0), Cm(3.0)])
add_source(doc, "Source: Author's analysis of Malta_.xlsx (sheet Reliability_Report).")

add_heading(doc, "3.7.3 Construct Validity - Principal Component Analysis (PCA)", level=2)
add_para(doc,
    "KMO measure of sampling adequacy was 0.91 (> 0.80 threshold) and Bartlett's test of "
    "sphericity was significant (χ² = 14,732, df = 171, p < 0.001), confirming factorability "
    "(Kaiser, 1974; Field, 2018). For the Farm-stage block (illustrative), three components with "
    "eigenvalues > 1 were extracted, jointly explaining 68.0% of variance (Table 3.5).")
add_caption(doc, "Table 3.5  PCA Solution for the Farm-Stage Block (Illustrative)")
add_table(doc,
    headers=["Component", "% Variance explained", "Items loading", "Interpretation", "Mean loading"],
    rows=[
        ["PC1", "27.2%", "7", "Cost factor", "0.510"],
        ["PC2", "21.6%", "6", "Quality factor", "0.410"],
        ["PC3", "19.2%", "6", "Time factor", "0.384"],
        ["Cumulative (3 PCs)", "68.0%", "19", "—", "0.435"],
    ],
    col_widths=[Cm(3.0), Cm(3.5), Cm(2.5), Cm(3.5), Cm(2.5)])
add_source(doc, "Source: Malta_.xlsx (PCA_Summary).")

add_heading(doc, "3.7.4 Validity Triangulation", level=2)
for bullet in [
    "Content validity - Lawshe CVR (Section 3.6.2).",
    "Construct validity - KMO/Bartlett + PCA (Section 3.7.3).",
    "Convergent validity - high mean inter-item correlations within sub-scales (Table 3.4) "
    "(Fornell & Larcker, 1981).",
    "Face validity - pilot expert review of language and logic.",
]:
    p = doc.add_paragraph(bullet, style="List Bullet")
    style_paragraph(p, space_after=4)

# ============================================================================
# 3.8 Analytical Framework
# ============================================================================
add_heading(doc, "3.8 Analytical Framework: Fuzzy AHP and Fuzzy ELECTRE", level=1)
add_para(doc,
    "The 82 validated factors were analysed using a two-stage Fuzzy MCDM framework: Fuzzy AHP for "
    "criterion weighting and Fuzzy ELECTRE for outranking of distribution-channel alternatives. "
    "Fuzzy logic is preferred over crisp methods because stakeholder evaluations are inherently "
    "linguistic and uncertain (Zadeh, 1965; Bellman & Zadeh, 1970).")

add_heading(doc, "3.8.1 Linguistic Variables and Triangular Fuzzy Numbers (TFNs)", level=2)
add_para(doc,
    "Linguistic ratings obtained from experts and channel members were mapped to TFNs using "
    "Chang's (1996) scale.")
add_caption(doc, "Table 3.6  Linguistic Scale and Corresponding Triangular Fuzzy Numbers")
add_table(doc,
    headers=["Linguistic term", "TFN (l, m, u)", "Reciprocal TFN"],
    rows=[
        ["Equally important", "(1, 1, 1)", "(1, 1, 1)"],
        ["Weakly important", "(2, 3, 4)", "(1/4, 1/3, 1/2)"],
        ["Fairly important", "(4, 5, 6)", "(1/6, 1/5, 1/4)"],
        ["Strongly important", "(6, 7, 8)", "(1/8, 1/7, 1/6)"],
        ["Absolutely important", "(8, 9, 9)", "(1/9, 1/9, 1/8)"],
    ],
    col_widths=[Cm(5.0), Cm(4.0), Cm(4.0)])
add_source(doc, "Source: Chang (1996).")

add_heading(doc, "3.8.2 Fuzzy Analytic Hierarchy Process (Fuzzy AHP)", level=2)
add_para(doc,
    "Following Chang's (1996) extent-analysis method, the procedure is:")
for step in [
    "Step 1 - Construct the fuzzy pairwise-comparison matrix Ã = [ã_ij]_{n × n}, where each "
    "ã_ij = (l_ij, m_ij, u_ij).",
    "Step 2 - Compute the fuzzy synthetic extent of the i-th object: "
    "S_i = (sum over j of M̃_{g_i}^{j}) ⊗ [sum over i and j of M̃_{g_i}^{j}]^(-1).",
    "Step 3 - Degree of possibility V(S2 ≥ S1): equals 1 if m2 ≥ m1; equals 0 if l1 ≥ u2; "
    "otherwise (l1 - u2) / [(m2 - u2) - (m1 - l1)].",
    "Step 4 - Compute the priority weight vector W = (d'(A1), d'(A2), ..., d'(An))^T where "
    "d'(A_i) = min V(S_i ≥ S_k), k ≠ i. Normalise W to obtain non-fuzzy priority weights.",
    "Step 5 - Consistency check using the Consistency Ratio computed on the defuzzified matrix; "
    "CR ≤ 0.10 is required (Saaty, 1980).",
]:
    p = doc.add_paragraph(step)
    style_paragraph(p, space_after=4)
add_para(doc, "The hierarchy implemented in this study is:")
for bullet in [
    "Goal (Level 0): Optimisation of Malta distribution channels.",
    "Main Criteria (Level 1): Cost, Time, Quality.",
    "Sub-criteria (Level 2): the five supply-chain stages.",
    "Factors (Level 3): the 82 validated items.",
]:
    p = doc.add_paragraph(bullet, style="List Bullet")
    style_paragraph(p, space_after=4)

add_heading(doc, "3.8.3 Fuzzy ELECTRE", level=2)
add_para(doc,
    "The ELECTRE family uses outranking relations rather than utility maximisation (Roy, 1991). "
    "The fuzzy variant adapted from Hatami-Marbini and Tavana (2011) and Sevkli (2010) was "
    "implemented in seven steps:")
for i, step in enumerate([
    "Decision matrix X̃ = [x̃_ij]_{m × n}, where rows are alternative distribution channels and "
    "columns are the fuzzy-AHP-weighted criteria.",
    "Normalisation of the fuzzy decision matrix.",
    "Weighted normalised matrix Ṽ = w̃_j ⊗ r̃_ij.",
    "Concordance set C_kl = { j | ṽ_kj ≽ ṽ_lj } and discordance set D_kl.",
    "Concordance and discordance indices C_kl, D_kl.",
    "Concordance dominance matrix F (using threshold c̄) and discordance dominance matrix G "
    "(using threshold d̄).",
    "Aggregate dominance matrix E = F · G; alternatives are ranked by net dominance score.",
], start=1):
    p = doc.add_paragraph(f"{i}. {step}")
    style_paragraph(p, space_after=4)
add_para(doc, "Three candidate distribution-channel configurations were evaluated:")
for bullet in [
    "A1: Existing multi-tier channel (Farmer -> Village trader -> Commission agent -> Wholesaler "
    "-> Retailer -> Consumer).",
    "A2: Cooperative-led short channel (Farmer -> FPO -> Wholesaler/Retailer -> Consumer).",
    "A3: Direct-marketing / e-platform channel (Farmer -> Aggregator/E-Mandi -> Consumer).",
]:
    p = doc.add_paragraph(bullet, style="List Bullet")
    style_paragraph(p, space_after=4)

add_heading(doc, "3.8.4 Software and Computation", level=2)
for bullet in [
    "Quantitative analysis (descriptive, reliability, PCA): IBM SPSS v.27 and Jamovi v.2.4.",
    "Fuzzy AHP and Fuzzy ELECTRE: custom MS-Excel workbook with VBA macros, cross-validated using "
    "Python 3.11 (numpy, pandas, scikit-criteria).",
    "Qualitative interview coding: NVivo v.14 (thematic analysis following Braun & Clarke, 2006).",
]:
    p = doc.add_paragraph(bullet, style="List Bullet")
    style_paragraph(p, space_after=4)

# ============================================================================
# 3.9 Validity, Reliability, Ethics
# ============================================================================
add_heading(doc, "3.9 Validity, Reliability, and Ethical Considerations", level=1)

add_heading(doc, "3.9.1 Validity", level=2)
add_para(doc,
    "Content, construct, convergent, and face validity were established as documented in Sections 3.6 and 3.7.")

add_heading(doc, "3.9.2 Reliability", level=2)
add_para(doc,
    "Internal-consistency reliability is reported in Table 3.4. Test-retest reliability was "
    "estimated on a sub-sample of 40 respondents at a two-week interval (r = 0.84). Inter-rater "
    "reliability for qualitative coding was Cohen's κ = 0.81 (Landis & Koch, 1977).")

add_heading(doc, "3.9.3 Trustworthiness of the Qualitative Strand", level=2)
add_para(doc,
    "Following Lincoln and Guba (1985), credibility was ensured by triangulation; transferability "
    "by thick description of context; dependability by an audit trail; and confirmability by "
    "member checking with five experts.")

add_heading(doc, "3.9.4 Ethical Considerations", level=2)
add_para(doc,
    "The study followed the ethical framework of the Indian Council of Social Science Research "
    "(ICSSR, 2020). Informed consent was obtained, anonymity and confidentiality were guaranteed, "
    "participation was voluntary, and the research protocol was reviewed by the Institutional "
    "Research Committee.")

# ============================================================================
# 3.10 Limitations
# ============================================================================
add_heading(doc, "3.10 Limitations of the Methodology", level=1)
for i, item in enumerate([
    "Geographical scope - restricted to four Garhwal districts; Kumaon and other citrus belts excluded.",
    "Cross-sectional snapshot - captures one season; does not model inter-annual variation.",
    "Self-reported perceptions - vulnerable to social-desirability bias, mitigated but not eliminated.",
    "Subjectivity of fuzzy linguistic mappings - partially controlled through expert calibration.",
    "Three channel alternatives - exhaustive options (e.g., processor-led channels) were excluded for tractability.",
], start=1):
    p = doc.add_paragraph(f"{i}. {item}")
    style_paragraph(p, space_after=4)

# ============================================================================
# 3.11 Summary
# ============================================================================
add_heading(doc, "3.11 Chapter Summary", level=1)
add_para(doc,
    "This chapter has set out the cross-sectional, mixed-methods, fuzzy-MCDM-based research "
    "design used to identify and prioritise the factors influencing optimisation of Malta "
    "distribution channels in Garhwal. From an initial pool of 144 candidate factors generated "
    "through literature review, field observation, and expert consultation, a rigorous "
    "CVR-reliability-PCA pipeline distilled a final set of 82 statistically validated factors "
    "organised across five supply-chain stages and three evaluative dimensions (Cost, Time, "
    "Quality). Data were collected from 1,300 stakeholders across the four major Malta-growing "
    "districts and analysed through Fuzzy AHP for weight estimation and Fuzzy ELECTRE for channel "
    "ranking. The reliability evidence (α ≥ 0.82 across all sub-scales) and the validity "
    "triangulation (CVR, KMO/Bartlett, PCA) lend high credibility to the analytical foundations "
    "of the next chapter.")

# ============================================================================
# References
# ============================================================================
doc.add_page_break()
add_heading(doc, "REFERENCES (Chapter 3)", level=1)

references = [
    "Bellman, R. E., & Zadeh, L. A. (1970). Decision-making in a fuzzy environment. Management Science, 17(4), B-141-B-164.",
    "Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in Psychology, 3(2), 77-101.",
    "Brislin, R. W. (1970). Back-translation for cross-cultural research. Journal of Cross-Cultural Psychology, 1(3), 185-216.",
    "Bryman, A., & Bell, E. (2015). Business Research Methods (4th ed.). Oxford University Press.",
    "Chang, D. Y. (1996). Applications of the extent analysis method on fuzzy AHP. European Journal of Operational Research, 95(3), 649-655.",
    "Chauhan, R. S., & Singh, V. (2018). Marketing of citrus fruits in Uttarakhand: an economic analysis. Indian Journal of Agricultural Marketing, 32(2), 101-115.",
    "Cochran, W. G. (1977). Sampling Techniques (3rd ed.). John Wiley & Sons.",
    "Connelly, L. M. (2008). Pilot studies. Medsurg Nursing, 17(6), 411-412.",
    "Creswell, J. W., & Creswell, J. D. (2018). Research Design: Qualitative, Quantitative, and Mixed Methods Approaches (5th ed.). SAGE.",
    "Creswell, J. W., & Plano Clark, V. L. (2017). Designing and Conducting Mixed Methods Research (3rd ed.). SAGE.",
    "DeVellis, R. F. (2017). Scale Development: Theory and Applications (4th ed.). SAGE.",
    "Directorate of Economics & Statistics, Uttarakhand. (2022). Statistical Diary of Uttarakhand. Government of Uttarakhand.",
    "Directorate of Horticulture, Uttarakhand. (2022). Annual Horticulture Report. Government of Uttarakhand.",
    "Etikan, I., Musa, S. A., & Alkassim, R. S. (2016). Comparison of convenience and purposive sampling. American Journal of Theoretical and Applied Statistics, 5(1), 1-4.",
    "Field, A. (2018). Discovering Statistics Using IBM SPSS Statistics (5th ed.). SAGE.",
    "Fornell, C., & Larcker, D. F. (1981). Evaluating structural equation models with unobservable variables and measurement error. Journal of Marketing Research, 18(1), 39-50.",
    "Government of Uttarakhand. (2022). State Profile. Dehradun: Information & Public Relations Department.",
    "Govindan, K., Rajendran, S., Sarkis, J., & Murugesan, P. (2015). Multi-criteria decision making approaches for green supplier evaluation and selection: A literature review. Journal of Cleaner Production, 98, 66-83.",
    "Gunasekaran, A., Patel, C., & McGaughey, R. E. (2004). A framework for supply chain performance measurement. International Journal of Production Economics, 87(3), 333-347.",
    "Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2019). Multivariate Data Analysis (8th ed.). Cengage.",
    "HAPPRC. (2021). Annual Report of the High Altitude Plant Physiology Research Centre. HNB Garhwal University, Srinagar.",
    "Hatami-Marbini, A., & Tavana, M. (2011). An extension of the ELECTRE I method for group decision-making under a fuzzy environment. Omega, 39(4), 373-386.",
    "ICAR-CITH. (2021). Vision 2030: Central Institute of Temperate Horticulture. ICAR.",
    "ICSSR. (2020). Ethical Guidelines for Social Science Research. Indian Council of Social Science Research.",
    "Kahraman, C., Cebeci, U., & Ulukan, Z. (2003). Multi-criteria supplier selection using fuzzy AHP. Logistics Information Management, 16(6), 382-394.",
    "Kaiser, H. F. (1974). An index of factorial simplicity. Psychometrika, 39(1), 31-36.",
    "Kothari, C. R. (2004). Research Methodology: Methods and Techniques (2nd rev. ed.). New Age International.",
    "Kumar, P., Singh, R., & Bhatt, D. (2019). Status and prospects of Malta cultivation in Uttarakhand hills. Indian Horticulture, 64(3), 18-22.",
    "Kvale, S., & Brinkmann, S. (2015). InterViews: Learning the Craft of Qualitative Research Interviewing (3rd ed.). SAGE.",
    "Landis, J. R., & Koch, G. G. (1977). The measurement of observer agreement for categorical data. Biometrics, 33(1), 159-174.",
    "Lawshe, C. H. (1975). A quantitative approach to content validity. Personnel Psychology, 28(4), 563-575.",
    "Levin, K. A. (2006). Study design III: Cross-sectional studies. Evidence-Based Dentistry, 7(1), 24-25.",
    "Lincoln, Y. S., & Guba, E. G. (1985). Naturalistic Inquiry. SAGE.",
    "Malhotra, N. K., & Dash, S. (2016). Marketing Research: An Applied Orientation (7th ed.). Pearson.",
    "Mentzer, J. T., DeWitt, W., Keebler, J. S., Min, S., Nix, N. W., Smith, C. D., & Zacharia, Z. G. (2001). Defining supply chain management. Journal of Business Logistics, 22(2), 1-25.",
    "Negi, S., & Anand, N. (2015). Supply chain of fruits and vegetables' agribusiness in Uttarakhand: Major issues and challenges. Journal of Supply Chain Management Systems, 4(1-2), 43-57.",
    "NHB. (2023). Indian Horticulture Database 2022-23. National Horticulture Board, Gurugram.",
    "Nunnally, J. C., & Bernstein, I. H. (1994). Psychometric Theory (3rd ed.). McGraw-Hill.",
    "Pant, K. S., Sharma, R., & Negi, V. (2020). Production constraints of Malta in mid-Himalayan region. Indian Journal of Hill Farming, 33(1), 56-63.",
    "Patton, M. Q. (2015). Qualitative Research and Evaluation Methods (4th ed.). SAGE.",
    "Punjabi, M. (2009). India: Increasing demand challenges the dairy and horticulture sectors. In Agro-industries for Development (pp. 192-217). FAO/CABI.",
    "Roy, B. (1991). The outranking approach and the foundations of ELECTRE methods. Theory and Decision, 31(1), 49-73.",
    "Saaty, T. L. (1980). The Analytic Hierarchy Process. McGraw-Hill.",
    "Saunders, M., Lewis, P., & Thornhill, A. (2019). Research Methods for Business Students (8th ed.). Pearson.",
    "Sedgwick, P. (2014). Cross-sectional studies: Advantages and disadvantages. BMJ, 348, g2276.",
    "Sevkli, M. (2010). An application of the fuzzy ELECTRE method for supplier selection. International Journal of Production Research, 48(12), 3393-3405.",
    "Singh, A., & Bhatt, D. (2018). Post-harvest losses in Malta in Uttarakhand. Progressive Horticulture, 50(1), 71-76.",
    "Tabachnick, B. G., & Fidell, L. S. (2019). Using Multivariate Statistics (7th ed.). Pearson.",
    "Tashakkori, A., & Teddlie, C. (2010). SAGE Handbook of Mixed Methods in Social & Behavioral Research (2nd ed.). SAGE.",
    "Yin, R. K. (2018). Case Study Research and Applications: Design and Methods (6th ed.). SAGE.",
    "Zadeh, L. A. (1965). Fuzzy sets. Information and Control, 8(3), 338-353.",
]
for ref in references:
    add_reference(doc, ref)

out = "/projects/sandbox/PP/Chapter_3_Research_Methodology.docx"
doc.save(out)
print(f"Saved: {out}")
